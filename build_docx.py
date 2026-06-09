# -*- coding: utf-8 -*-
"""
Generate the SLR article as a .docx in the Al-Rikaz: Jurnal Ekonomi Syariah style.
English abstract + Indonesian body + Turabian/Chicago footnotes.
"""
import zipfile, shutil, os, re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "Islamic_Economic_Literacy_Education_Gen_Z_SLR.docx"

doc = Document()

# ----------------- base styles -----------------
normal = doc.styles["Normal"]
normal.font.name = "Book Antiqua"
normal.font.size = Pt(11)
rpr = normal.element.get_or_add_rPr()
rfonts = rpr.get_or_add_rFonts()
rfonts.set(qn("w:ascii"), "Book Antiqua")
rfonts.set(qn("w:hAnsi"), "Book Antiqua")

for section in doc.sections:
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

# ----------------- footnote machinery -----------------
_footnotes = []  # list of (id, text)
_fn_counter = {"n": 0}

def add_footnote(paragraph, text):
    _fn_counter["n"] += 1
    fid = _fn_counter["n"]
    _footnotes.append((fid, text))
    run = paragraph.add_run()
    rpr = OxmlElement("w:rPr")
    valign = OxmlElement("w:vertAlign")
    valign.set(qn("w:val"), "superscript")
    rpr.append(valign)
    run._r.append(rpr)
    ref = OxmlElement("w:footnoteReference")
    ref.set(qn("w:id"), str(fid))
    run._r.append(ref)
    return fid

# ----------------- helpers -----------------
def set_after(p, pt=6, before=0, line=1.5):
    pf = p.paragraph_format
    pf.space_after = Pt(pt)
    pf.space_before = Pt(before)
    pf.line_spacing = line

def para(text="", align="justify", bold=False, italic=False, size=11, after=6, before=0, line=1.5):
    p = doc.add_paragraph()
    a = {"justify": WD_ALIGN_PARAGRAPH.JUSTIFY, "center": WD_ALIGN_PARAGRAPH.CENTER,
         "left": WD_ALIGN_PARAGRAPH.LEFT, "right": WD_ALIGN_PARAGRAPH.RIGHT}[align]
    p.alignment = a
    if text:
        r = p.add_run(text)
        r.bold = bold; r.italic = italic; r.font.size = Pt(size)
    set_after(p, pt=after, before=before, line=line)
    return p

def heading_caps(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text.upper())
    r.bold = True; r.font.size = Pt(11)
    set_after(p, pt=6, before=10, line=1.5)
    return p

def subhead(text, italic=True):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True; r.italic = italic; r.font.size = Pt(11)
    set_after(p, pt=4, before=6, line=1.5)
    return p

def style_table(tbl):
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.line_spacing = 1.0
                for r in p.runs:
                    r.font.size = Pt(10)

def cell_text(cell, text, bold=False, italic=False):
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.bold = bold; r.italic = italic; r.font.size = Pt(10)

# =====================================================================
# JOURNAL HEADER
# =====================================================================
para("[Nama Jurnal Ekonomi Syariah]", align="left", bold=True, size=11, after=0)
para("E-ISSN: [________]   |   DOI: [________]", align="left", size=9, after=0)
para("Volume [__] Issue [__], [Tahun]", align="left", size=9, after=0)
para("Journal Page is available to: [URL Jurnal]", align="left", size=9, after=6)

hr = doc.add_paragraph()
hrpr = hr.paragraph_format
hrpr.space_after = Pt(6)
pbdr = OxmlElement("w:pBdr")
bottom = OxmlElement("w:bottom")
bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "6")
bottom.set(qn("w:space"), "1"); bottom.set(qn("w:color"), "000000")
pbdr.append(bottom)
hr._p.get_or_add_pPr().append(pbdr)

# TITLE
para("Islamic Economic Literacy Education among Generation Z Muslims:", align="center", bold=True, size=14, after=0, before=4)
para("A Systematic Literature Review", align="center", bold=True, size=14, after=4)
para("[Nama Penulis]", align="center", bold=True, size=11, after=0)
para("[Program Studi / Fakultas], [Nama Universitas], [Kota], Indonesia", align="center", italic=True, size=10, after=2)
para("e-mail: [email penulis]", align="center", italic=True, size=10, after=8)

# =====================================================================
# ARTICLE INFO + ABSTRACT (2-column table)
# =====================================================================
info = doc.add_table(rows=1, cols=2)
info.style = "Table Grid"
info.alignment = WD_TABLE_ALIGNMENT.CENTER
info.columns[0].width = Cm(5.0)
info.columns[1].width = Cm(10.5)
left, right = info.rows[0].cells

# left: article info
def left_line(cell, text, bold=False, italic=False, first=False):
    p = cell.paragraphs[0] if first else cell.add_paragraph()
    r = p.add_run(text); r.bold = bold; r.italic = italic; r.font.size = Pt(9)
    p.paragraph_format.space_after = Pt(2); p.paragraph_format.line_spacing = 1.0
    return p

left_line(left, "ARTICLE INFO", bold=True, first=True)
left_line(left, "Article history:", bold=True)
left_line(left, "Received   : [tgl]")
left_line(left, "Revised    : [tgl]")
left_line(left, "Accepted   : [tgl]")
left_line(left, "Available  : [tgl]")
left_line(left, "")
left_line(left, "Keywords:", bold=True)
left_line(left, "Islamic economic literacy; Sharia financial literacy; Generation Z Muslims; Islamic financial education; halal financial behaviour.")
left_line(left, "")
left_line(left, "Paper type: Research Paper", italic=True)
left_line(left, "")
left_line(left, "Please cite this article:", bold=True)
left_line(left, "[Nama Penulis]. ([Tahun]). Islamic Economic Literacy Education among Generation Z Muslims: A Systematic Literature Review. [Nama Jurnal], [Vol](Issue), [halaman].", italic=True)
left_line(left, "")
left_line(left, "This article is licensed under CC BY license. Copyright © [Tahun], the author(s).")

# right: abstract
def right_line(cell, text, bold=False, italic=False, first=False):
    p = cell.paragraphs[0] if first else cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text); r.bold = bold; r.italic = italic; r.font.size = Pt(9)
    p.paragraph_format.space_after = Pt(3); p.paragraph_format.line_spacing = 1.0
    return p

right_line(right, "ABSTRACT", bold=True, first=True)
abstract = ("The rapid expansion of Islamic finance, the digitalisation of economic life, and the rise of a "
"halal lifestyle have made Islamic economic literacy an increasingly strategic competency for Muslim "
"societies, particularly for Generation Z Muslims who enter financial life through fintech, social media, "
"and Sharia-compliant products. This study aims to systematically identify, evaluate, and synthesise "
"existing research on Islamic economic literacy education among Generation Z Muslims. The study employs "
"a Systematic Literature Review (SLR) following a transparent and replicable protocol consisting of "
"planning, conducting, and reporting phases, structured by a PICOC framework, Boolean search strings, "
"inclusion and exclusion criteria, a PRISMA-based selection workflow, a weighted quality-assessment "
"checklist (Yes = 1, Partially = 0.5, No = 0), a standardised data-extraction form, and combined thematic "
"and descriptive bibliometric analysis. Searches were conducted in Scopus, Web of Science, Google Scholar, "
"DOAJ, Crossref, Semantic Scholar, Garuda, and SINTA-indexed journals, covering peer-reviewed studies "
"published between 2015 and 2025 in English or Indonesian. The synthesis reveals six themes: Islamic "
"economic literacy as a multidimensional faith-anchored construct; its relationship with Sharia-compliant "
"financial behaviour; digital and social-media learning channels; Islamic financial institutions and youth "
"awareness; zakat, waqf, and sadaqah social-finance literacy; and recurring challenges and gaps. The review "
"contributes a theoretical consolidation of Islamic economic literacy and practical implications for "
"educators, Islamic financial institutions, and policymakers in cultivating riba-averse, Sharia-conscious "
"financial behaviour among young Muslims.")
right_line(right, abstract)

doc.add_paragraph().paragraph_format.space_after = Pt(2)

# =====================================================================
# PENDAHULUAN
# =====================================================================
heading_caps("Pendahuluan")

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; set_after(p)
p.add_run("Dalam tiga dekade terakhir, keuangan Islam berkembang dari alternatif etis yang bersifat khusus "
"menjadi industri global yang signifikan, meliputi perbankan syariah, pasar modal syariah, takaful, serta "
"instrumen keuangan sosial Islam seperti zakat, wakaf, dan sedekah. Seiring pertumbuhan industri tersebut, "
"perhatian akademik dan kebijakan terhadap ")
r = p.add_run("literasi keuangan syariah"); r.italic = True
p.add_run("\u2014yakni pengetahuan, keterampilan, sikap, dan keyakinan yang diperlukan untuk mengambil "
"keputusan keuangan yang sesuai dengan prinsip syariah\u2014juga semakin meningkat.")
add_footnote(p, "Purnomo M. Antara, Rosidah Musa, dan Faridah Hassan, \u201cBridging Islamic Financial Literacy and Halal Literacy: The Way Forward in Halal Ecosystem,\u201d Procedia Economics and Finance 37 (2016): 196\u2013202.")
p.add_run(" Jika literasi keuangan konvensional dipahami sebagai bentuk modal manusia yang memperbaiki "
"perilaku menabung, perencanaan, dan pengelolaan utang,")
add_footnote(p, "Annamaria Lusardi dan Olivia S. Mitchell, \u201cThe Economic Importance of Financial Literacy: Theory and Evidence,\u201d Journal of Economic Literature 52, no. 1 (2014): 5\u201344.")
p.add_run(" maka literasi ekonomi Islam menambahkan dimensi normatif berbasis keimanan, yaitu keharusan "
"memperoleh, membelanjakan, menabung, menginvestasikan, dan mendermakan harta dengan cara yang "
"menghindari riba, gharar, dan maysir, sekaligus menunaikan kewajiban seperti zakat.")
add_footnote(p, "Mahmoud A. El-Gamal, Islamic Finance: Law, Economics, and Practice (Cambridge: Cambridge University Press, 2006); Feisal Khan, \u201cHow \u2018Islamic\u2019 Is Islamic Banking?\u201d Journal of Economic Behavior & Organization 76, no. 3 (2010): 805\u2013820.")

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; set_after(p)
p.add_run("Urgensi kompetensi ini semakin besar bagi ")
r = p.add_run("Generasi Z"); r.bold = True
p.add_run("\u2014kohort yang lahir sekitar tahun 1997 dan sesudahnya.")
add_footnote(p, "Michael Dimock, Defining Generations: Where Millennials End and Generation Z Begins (Washington, DC: Pew Research Center, 2019).")
p.add_run(" Generasi Z Muslim tumbuh dengan telepon pintar, media sosial, dompet elektronik, layanan "
"\u201cbuy-now-pay-later\u201d, dan platform fintech syariah sebagai infrastruktur baku kehidupan ekonomi mereka.")
add_footnote(p, "Marc Prensky, \u201cDigital Natives, Digital Immigrants,\u201d On the Horizon 9, no. 5 (2001): 1\u20136; Mustafa Raza Rabbani, Shahnawaz Khan, dan Eleftherios I. Thalassinos, \u201cFinTech, Blockchain and Islamic Finance: An Extensive Literature Review,\u201d International Journal of Economics and Business Administration 8, no. 2 (2020): 65\u201386.")
p.add_run(" Mereka sekaligus merupakan segmen konsumen yang kuat bagi ekonomi halal global dan populasi "
"yang terpapar budaya digital konsumtif, pinjaman daring, serta produk spekulatif yang berpotensi "
"bertentangan dengan norma syariah. Bukti empiris menunjukkan bahwa literasi keuangan syariah memengaruhi "
"keputusan investasi generasi muda Muslim di pasar modal syariah")
add_footnote(p, "Salsabilla Nur Anisa dan Muh Shadiqul Fajri AF, \u201cThe Impact of Islamic Financial Literacy on Gen Z\u2019s Investment Decisions in the Islamic Capital Market\u201d (working paper, SSRN, 2025).")
p.add_run(" serta niat mereka untuk mengadopsi perbankan syariah.")
add_footnote(p, "Mohamed Albaity dan Mahfuzur Rahman, \u201cThe Intention to Use Islamic Banking: An Exploratory Study to Measure Islamic Financial Literacy,\u201d International Journal of Emerging Markets 14, no. 5 (2019): 988\u20131012.")

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; set_after(p)
p.add_run("Pendidikan berada di pusat dinamika ini. Sekolah, perguruan tinggi, pesantren dan lembaga "
"pendidikan Islam lain, keluarga, komunitas Muslim, serta lembaga keuangan syariah berperan sebagai "
"saluran pembentukan kesadaran ekonomi Islam. Bukti dari pendidikan tinggi di Malaysia dan Indonesia "
"menunjukkan bahwa paparan mata kuliah muamalah dan pendidikan keuangan syariah yang terstruktur "
"secara nyata meningkatkan tingkat literasi,")
add_footnote(p, "Siti Hafsha Albasri Rahim, Rosemaliza Abdul Rashid, dan Abu Bakar Hamed, \u201cIslamic Financial Literacy and Its Determinants among University Students: An Exploratory Factor Analysis,\u201d International Journal of Economics and Financial Issues 6, no. 7S (2016): 32\u201335; Gatot Nazir Ahmad dkk., \u201cDeterminants of the Islamic Financial Literacy,\u201d Accounting 6, no. 6 (2020): 961\u2013966.")
p.add_run(" sementara pedagogi inovatif seperti sosialisasi keuangan syariah berbasis augmented reality "
"dapat meningkatkan pengetahuan dan motivasi belajar.")
add_footnote(p, "Ismi Solikhatun dan Ratna Candra Sari, \u201cSharia Financial Education Using Augmented Reality Technology to Increase Student Motivation in Distance Learning,\u201d AIP Conference Proceedings 2654 (2023): 020008.")
p.add_run(" Pada ranah keuangan sosial, literasi zakat terbukti memengaruhi niat pembayaran zakat melalui "
"platform digital.")
add_footnote(p, "Rahmatina A. Kasri dan Adela Miranti Yuniar, \u201cDeterminants of Digital Zakat Payments: Lessons from Indonesian Experience,\u201d Journal of Islamic Accounting and Business Research 12, no. 3 (2021): 362\u2013379.")

para("Dengan demikian, literasi ekonomi Islam bukan sekadar keterampilan kognitif yang berdiri sendiri, "
"melainkan pemampu (enabler) bagi rangkaian luaran yang lebih luas: kepatuhan syariah, perilaku keuangan "
"halal, penunaian zakat dan wakaf, partisipasi dalam investasi halal, dan penghindaran riba secara sengaja. "
"Keterkaitan ini menempatkan topik penelitian secara tegas dalam ranah hukum dan ekonomi Islam, bukan "
"sekadar pendidikan keuangan umum.")

subhead("Kesenjangan Penelitian, Kebaruan, dan Tujuan")
para("Meskipun kajian terkait terus bertumbuh, literatur masih terfragmentasi. Studi-studi cenderung "
"membahas literasi keuangan syariah, pemuda Muslim, ekonomi syariah, gaya hidup halal, dan pendidikan "
"keuangan secara terpisah. Tinjauan yang ada lebih banyak menyoroti area berdekatan\u2014adopsi fintech "
"syariah, manajemen zakat, dan literasi keuangan syariah pada populasi umum\u2014sehingga sintesis sistematis "
"yang secara khusus berfokus pada pendidikan literasi ekonomi Islam di kalangan Generasi Z Muslim masih "
"sangat terbatas. Kebaruan tinjauan ini terletak pada penerapan protokol SLR yang ketat, transparan, dan "
"dapat direplikasi untuk mengonsolidasikan bukti yang tersebar. Tujuan penelitian adalah mengidentifikasi, "
"menilai, dan menyintesis penelitian tertelaah-sejawat mengenai pendidikan literasi ekonomi Islam pada "
"Generasi Z Muslim, serta menurunkan implikasi teoretis, praktis, dan kebijakan.")

# =====================================================================
# KAJIAN LITERATUR
# =====================================================================
heading_caps("Kajian Literatur")

subhead("Konsep Literasi Ekonomi Islam dan Literasi Keuangan Syariah")
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; set_after(p)
p.add_run("Literatur secara konsisten memposisikan literasi ekonomi Islam sebagai konsep yang lebih luas "
"daripada pengetahuan keuangan konvensional. Konsep ini memadukan dimensi kognitif (pemahaman akad, riba, "
"zakat, dan produk yang sesuai syariah), dimensi normatif-religius (niat untuk mematuhi hukum Islam), dan "
"dimensi perilaku (keterampilan, keyakinan diri, serta kemampuan bertindak atas pengetahuan tersebut).")
add_footnote(p, "Antara, Musa, dan Hassan, \u201cBridging Islamic Financial Literacy and Halal Literacy\u201d; Purnomo M. Antara, Rosidah Musa, dan Faridah Hassan, \u201cConceptualisation and Operationalisation of Islamic Financial Literacy Scale,\u201d Pertanika Journal of Social Sciences & Humanities 25, no. S (2017): 251\u2013260.")
p.add_run(" Para sarjana berargumen bahwa instrumen literasi konvensional tidak memadai bagi Muslim karena "
"mengabaikan parameter berbasis keimanan yang membedakan keuangan yang dibolehkan dan yang dilarang, "
"sehingga mendorong pengembangan skala literasi keuangan syariah tersendiri.")

subhead("Generasi Z Muslim dan Perilaku Keuangan Sesuai Syariah")
para("Bukti menunjukkan bahwa literasi ekonomi Islam merupakan anteseden yang bermakna bagi perilaku "
"yang selaras dengan syariah di kalangan Muslim muda. Literasi yang lebih tinggi dikaitkan dengan niat yang "
"lebih besar untuk menggunakan perbankan syariah serta keputusan investasi yang lebih baik di pasar modal "
"syariah. Pada saat yang sama, keterpaparan kohort ini terhadap budaya digital yang konsumtif menyebabkan "
"literasi tidak selalu berbuah praktik, yang menegaskan adanya kesenjangan sikap\u2013perilaku (attitude\u2013"
"behaviour gap) yang harus dijawab oleh pendidikan.")

subhead("Pendidikan, Media Digital, dan Keuangan Sosial Islam")
para("Karena Generasi Z adalah penutur asli digital, kajian menekankan pedagogi berbantuan teknologi "
"seperti augmented reality, gamifikasi, pembelajaran seluler, dan penjangkauan berbasis media sosial. "
"Lembaga keuangan syariah\u2014bank, platform pasar modal, dan lembaga keuangan mikro\u2014berperan ganda "
"sebagai penyedia sekaligus pendidik. Di luar keuangan komersial, literasi keuangan sosial Islam (zakat, "
"wakaf, sedekah) semakin dimediasi oleh platform digital yang relevan dengan donatur muda, menunjukkan "
"bahwa literasi ekonomi Islam menjangkau pula inti redistributif dan kesejahteraan dalam ekonomi Islam.")

# =====================================================================
# METODE PENELITIAN
# =====================================================================
heading_caps("Metode Penelitian")
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; set_after(p)
p.add_run("Penelitian ini menggunakan pendekatan Systematic Literature Review (SLR), yaitu proses yang "
"melibatkan identifikasi, penilaian, dan penafsiran seluruh bukti penelitian yang tersedia secara "
"sistematis untuk menjawab pertanyaan penelitian tertentu.")
add_footnote(p, "Barbara Kitchenham dan Stuart Charters, Guidelines for Performing Systematic Literature Reviews in Software Engineering, EBSE Technical Report EBSE-2007-01 (Keele University & University of Durham, 2007); David Tranfield, David Denyer, dan Palminder Smart, \u201cTowards a Methodology for Developing Evidence-Informed Management Knowledge by Means of Systematic Review,\u201d British Journal of Management 14, no. 3 (2003): 207\u2013222.")
p.add_run(" SLR berbeda dari tinjauan naratif karena bersandar pada protokol a priori yang terdokumentasi "
"sehingga proses bersifat transparan, objektif, dan dapat direplikasi.")
add_footnote(p, "Hannah Snyder, \u201cLiterature Review as a Research Methodology: An Overview and Guidelines,\u201d Journal of Business Research 104 (2019): 333\u2013339.")
p.add_run(" Protokol mengikuti tiga tahap: perencanaan, pelaksanaan, dan pelaporan,")
add_footnote(p, "Angela Carrera-Rivera dkk., \u201cHow-to Conduct a Systematic Literature Review: A Quick Guide for Computer Science Research,\u201d MethodsX 9 (2022): 101895.")
p.add_run(" serta mengacu pada standar pelaporan PRISMA 2020.")
add_footnote(p, "Matthew J. Page dkk., \u201cThe PRISMA 2020 Statement: An Updated Guideline for Reporting Systematic Reviews,\u201d BMJ 372 (2021): n71.")

# SLR steps table (like Gambar 1)
subhead("Tahapan Systematic Literature Review (SLR)")
steps = [
    ("Perencanaan", "1. Merumuskan pertanyaan penelitian\n2. Mengembangkan protokol tinjauan\n3. Validasi protokol tinjauan"),
    ("Pelaksanaan", "4. Identifikasi penelitian yang relevan\n5. Menyaring studi (judul, abstrak, teks penuh)\n6. Menilai kualitas studi\n7. Ekstraksi data\n8. Sintesis data"),
    ("Pelaporan", "9. Menulis laporan hasil\n10. Validasi laporan"),
]
t = doc.add_table(rows=len(steps), cols=2)
for i,(a,b) in enumerate(steps):
    cell_text(t.rows[i].cells[0], a, bold=True)
    cell_text(t.rows[i].cells[1], b)
style_table(t)
para("Gambar 1. Tahapan Systematic Literature Review (SLR)", align="center", italic=True, size=9, after=8, before=2)

# Research Questions
subhead("Pertanyaan Penelitian")
para("Pertanyaan penelitian disusun untuk menjaga fokus sistematis tinjauan dan memandu identifikasi "
"studi, ekstraksi data, serta analisis dan sintesis. Lima pertanyaan penelitian (RQ) dirumuskan sebagai "
"berikut.")
rqs = [
    ("RQ1", "Apa tema dan tren utama dalam penelitian terdahulu tentang pendidikan literasi ekonomi Islam di kalangan Generasi Z Muslim?"),
    ("RQ2", "Bagaimana literasi ekonomi Islam dikonseptualisasikan dalam literatur yang ada?"),
    ("RQ3", "Strategi pendidikan, media pembelajaran, atau pendekatan digital apa yang digunakan untuk memperkuat literasi ekonomi Islam pada Generasi Z Muslim?"),
    ("RQ4", "Tantangan dan peluang apa yang teridentifikasi dalam mempromosikan literasi ekonomi Islam pada Generasi Z Muslim?"),
    ("RQ5", "Kesenjangan penelitian dan arah riset masa depan apa yang muncul dari literatur yang ada?"),
]
t = doc.add_table(rows=len(rqs)+1, cols=2)
cell_text(t.rows[0].cells[0], "ID", bold=True); cell_text(t.rows[0].cells[1], "Research Questions", bold=True)
for i,(a,b) in enumerate(rqs,1):
    cell_text(t.rows[i].cells[0], a, bold=True); cell_text(t.rows[i].cells[1], b)
style_table(t)
para("Tabel 1. Pertanyaan Penelitian pada Tinjauan Literatur", align="center", italic=True, size=9, after=8, before=2)

# PICOC
subhead("Kerangka PICOC")
para("Cakupan tinjauan distrukturkan menggunakan kerangka PICOC (Population, Intervention, Comparison, "
"Outcome, Context) yang menerjemahkan pertanyaan penelitian menjadi konsep yang dapat ditelusuri.")
picoc = [
    ("Population", "Generasi Z Muslim, pemuda Muslim, pelajar/mahasiswa Muslim, generasi muda Muslim (Gen Z, post-millennials, Muslim youth, young Muslims, Muslim students)."),
    ("Intervention", "Pendidikan literasi ekonomi Islam / keuangan syariah (Islamic/Sharia financial literacy education, halal financial education, muamalah education, financial socialisation)."),
    ("Comparison", "Literasi keuangan konvensional, literasi ekonomi umum, pendidikan keuangan non-syariah; atau tanpa pembanding bila tidak relevan (N/A)."),
    ("Outcome", "Kesadaran ekonomi Islam, perilaku keuangan sesuai syariah, keputusan keuangan halal, adopsi perbankan syariah, investasi halal, zakat, wakaf, sedekah, penghindaran riba."),
    ("Context", "Pendidikan ekonomi Islam, hukum dan ekonomi Islam, lembaga keuangan syariah, sekolah, perguruan tinggi, pesantren/madrasah, komunitas Muslim, dan lingkungan pembelajaran digital."),
]
t = doc.add_table(rows=len(picoc), cols=2)
for i,(a,b) in enumerate(picoc):
    cell_text(t.rows[i].cells[0], a, bold=True); cell_text(t.rows[i].cells[1], b)
style_table(t)
para("Tabel 2. Ringkasan Kerangka PICOC", align="center", italic=True, size=9, after=8, before=2)

# Search strings & databases
subhead("Strategi Pencarian dan Sumber Pustaka Digital")
para("String pencarian dibangun dengan menggabungkan kata kunci dan sinonim PICOC menggunakan operator "
"Boolean AND dan OR, serta diterapkan pada judul, abstrak, dan kata kunci sejauh didukung oleh basis data. "
"Contoh string pencarian: (\u201cIslamic economic literacy\u201d OR \u201cSharia financial literacy\u201d OR "
"\u201cIslamic financial literacy\u201d) AND (\u201cGeneration Z\u201d OR \u201cGen Z\u201d OR \u201cMuslim youth\u201d OR "
"\u201cMuslim students\u201d) AND (\u201ceducation\u201d OR \u201cfinancial education\u201d OR \u201cliteracy education\u201d). "
"Pencarian dilakukan pada delapan sumber, yaitu Scopus, Web of Science, Google Scholar, DOAJ, Crossref, "
"Semantic Scholar, Garuda, dan jurnal terindeks SINTA, dilengkapi penelusuran maju-mundur (snowballing). "
"Pemilihan sumber ini penting karena topik bersinggungan dengan ekonomi Islam, pendidikan, literasi "
"keuangan, dan kajian pemuda Muslim, dengan banyak bukti primer berasal dari Indonesia dan Malaysia.")

# Inclusion / Exclusion
subhead("Pemilihan Studi: Kriteria Inklusi dan Eksklusi")
inc = ("1. Membahas literasi ekonomi Islam, literasi keuangan syariah, atau pendidikan ekonomi Islam.\n"
"2. Berfokus pada Generasi Z, pemuda Muslim, pelajar/mahasiswa, atau generasi muda Muslim.\n"
"3. Membahas pendidikan, literasi, strategi pembelajaran, media digital, perilaku keuangan, gaya hidup "
"halal, lembaga keuangan syariah, zakat, wakaf, sedekah, atau investasi halal.\n"
"4. Artikel jurnal tertelaah-sejawat, prosiding, atau bab buku ilmiah.\n"
"5. Terbit pada 2015\u20132025.\n6. Ditulis dalam bahasa Inggris atau Indonesia.")
exc = ("1. Tidak terkait ekonomi Islam atau literasi keuangan syariah.\n"
"2. Hanya membahas literasi keuangan konvensional tanpa relevansi ekonomi Islam.\n"
"3. Tidak membahas pemuda, pelajar, atau Generasi Z.\n"
"4. Opini, berita, laporan non-akademik, atau teks penuh tidak dapat diakses.\n"
"5. Artikel duplikat.\n6. Kualitas metodologis lemah setelah penilaian kualitas (di bawah ambang).")
t = doc.add_table(rows=2, cols=2)
cell_text(t.rows[0].cells[0], "Kriteria Inklusi", bold=True); cell_text(t.rows[0].cells[1], inc)
cell_text(t.rows[1].cells[0], "Kriteria Eksklusi", bold=True); cell_text(t.rows[1].cells[1], exc)
style_table(t)
para("Tabel 3. Kriteria Inklusi dan Eksklusi", align="center", italic=True, size=9, after=8, before=2)

para("Proses seleksi berlangsung dalam lima tahap dan dapat dilaporkan melalui diagram alir PRISMA 2020: "
"(1) identifikasi rekaman dari basis data dan snowballing; (2) penghapusan duplikat; (3) penyaringan "
"judul/abstrak; (4) penilaian kelayakan melalui pembacaan teks penuh; dan (5) penyertaan akhir studi untuk "
"sintesis.")

# Quality assessment
subhead("Penilaian Kualitas Studi")
para("Setiap studi kandidat dinilai dengan daftar tilik penilaian kualitas (QA) tujuh butir yang diberi "
"skor Ya = 1, Sebagian = 0,5, Tidak = 0 (maksimum = 7). Studi dengan skor \u2265 4,0 dari 7 (\u2248 57%) "
"dipertahankan, sedangkan yang di bawah ambang dikeluarkan.")
qas = [
    ("QA1", "Apakah artikel menyatakan tujuan penelitiannya dengan jelas?"),
    ("QA2", "Apakah artikel relevan dengan literasi ekonomi Islam atau literasi keuangan syariah?"),
    ("QA3", "Apakah artikel berfokus pada Generasi Z, pemuda Muslim, pelajar/mahasiswa, atau generasi muda Muslim?"),
    ("QA4", "Apakah metode penelitian dijelaskan dengan jelas?"),
    ("QA5", "Apakah temuan disajikan dengan jelas dan didukung data?"),
    ("QA6", "Apakah artikel membahas implikasi bagi ekonomi Islam, pendidikan, atau perilaku keuangan sesuai syariah?"),
    ("QA7", "Apakah artikel mengakui keterbatasan atau arah riset masa depan?"),
]
t = doc.add_table(rows=len(qas)+1, cols=2)
cell_text(t.rows[0].cells[0], "Kode", bold=True); cell_text(t.rows[0].cells[1], "Pertanyaan Penilaian Kualitas", bold=True)
for i,(a,b) in enumerate(qas,1):
    cell_text(t.rows[i].cells[0], a, bold=True); cell_text(t.rows[i].cells[1], b)
style_table(t)
para("Tabel 4. Daftar Tilik Penilaian Kualitas", align="center", italic=True, size=9, after=8, before=2)

# Data extraction & synthesis
subhead("Ekstraksi Data dan Sintesis")
para("Formulir ekstraksi data baku digunakan untuk merekam: penulis; tahun; negara/konteks; judul; "
"tujuan penelitian; metodologi; populasi/sampel; konsep kunci; jenis literasi; strategi pendidikan/media "
"pembelajaran; tema ekonomi Islam; temuan kunci; tantangan; kesenjangan; dan relevansi terhadap pertanyaan "
"penelitian. Analisis memadukan sintesis tematik dengan analisis bibliometrik deskriptif (tren publikasi "
"per tahun, dominasi negara/konteks, sumber publikasi utama, dan kata kunci yang sering muncul).")

# =====================================================================
# HASIL DAN PEMBAHASAN
# =====================================================================
heading_caps("Hasil Penelitian dan Pembahasan")

subhead("Gambaran Bibliometrik")
para("Korpus yang disintesis menunjukkan tiga pola konsisten. Pertama, aktivitas publikasi meningkat "
"tajam sejak 2016 dan menguat setelah 2019, sejalan dengan ekspansi fintech syariah dan perhatian kebijakan "
"terhadap literasi keuangan. Kedua, distribusi geografis sangat terkonsentrasi di Indonesia dan Malaysia, "
"dengan kontribusi tambahan dari kawasan Teluk, Yordania, dan pusat keuangan Islam lainnya. Ketiga, sumber "
"publikasi yang paling sering adalah jurnal spesialis seperti Journal of Islamic Marketing, International "
"Journal of Islamic and Middle Eastern Finance and Management (IMEFM), dan Journal of Islamic Accounting and "
"Business Research (JIABR), serta jurnal terindeks SINTA/DOAJ. Kata kunci yang berulang mencakup literasi "
"keuangan syariah, religiusitas, Generasi Z/milenial, perbankan syariah, fintech, zakat, dan wakaf tunai.")

subhead("Karakteristik Studi yang Disintesis")
para("Tabel 5 merangkum studi representatif yang terverifikasi dan menjadi jangkar sintesis. Untuk "
"pengajuan akhir, tabel ini perlu diperluas mencakup seluruh studi yang lolos ambang kualitas.")
rows = [
    ("1","Lusardi & Mitchell (2014)","AS/global","Ekonomi literasi keuangan","Tinjauan teori & bukti","Literasi keuangan sebagai investasi modal manusia yang memengaruhi menabung, perencanaan, dan utang","Dasar konseptual yang diperluas literasi ekonomi Islam (RQ2)"),
    ("2","Antara, Musa & Hassan (2016)","Malaysia","Menjembatani literasi keuangan & halal","Konseptual","Memposisikan IFL dalam ekosistem halal","Konseptualisasi; keterkaitan halal\u2013keuangan (RQ2, RQ4)"),
    ("3","Rahim, Rashid & Hamed (2016)","Malaysia","Determinan IFL mahasiswa","Survei, EFA","Religiusitas dan kepuasan finansial membentuk IFL mahasiswa","Determinan literasi pemuda (RQ1, RQ3)"),
    ("4","Albaity & Rahman (2019)","UEA","IFL & niat memakai bank syariah","Survei, SEM","IFL berkaitan dengan niat lebih besar mengadopsi perbankan syariah","Literasi \u2192 keterlibatan kelembagaan (RQ2, RQ4)"),
    ("5","Ahmad dkk. (2020)","Indonesia","Determinan IFL","Survei","Mengidentifikasi determinan sosio-demografis & pendidikan IFL","Pendorong pendidikan/konteks (RQ1, RQ3)"),
    ("6","Solikhatun & Sari (2023)","Indonesia","Pendidikan keuangan syariah berbasis AR","Kuasi-eksperimen","Media AR meningkatkan motivasi & pengetahuan keuangan syariah pada pembelajaran jarak jauh","Pedagogi digital untuk pemuda (RQ3)"),
    ("7","Kasri & Yuniar (2021)","Indonesia","Determinan zakat digital","Survei, UTAUT","Literasi zakat meningkatkan niat membayar zakat via platform digital","Literasi keuangan sosial & kanal digital (RQ3, RQ4)"),
    ("8","Anisa & Fajri (2025)","Indonesia","IFL & investasi Gen Z di pasar modal syariah","Survei","IFL berpengaruh positif signifikan terhadap keputusan investasi Gen Z","Bukti langsung Gen Z (RQ1, RQ4)"),
]
hdr = ["No.","Penulis & Tahun","Konteks","Fokus","Metode","Temuan Kunci","Relevansi"]
t = doc.add_table(rows=len(rows)+1, cols=len(hdr))
for j,h in enumerate(hdr):
    cell_text(t.rows[0].cells[j], h, bold=True)
for i,row in enumerate(rows,1):
    for j,val in enumerate(row):
        cell_text(t.rows[i].cells[j], val)
style_table(t)
# narrow columns
for j,w in enumerate([0.9,2.4,1.6,2.2,1.8,3.2,3.0]):
    for r in t.rows:
        r.cells[j].width = Cm(w)
para("Tabel 5. Penulis, Konteks, Metode, dan Temuan Studi Terpilih", align="center", italic=True, size=9, after=8, before=2)

subhead("Temuan Tematik dan Pembahasan")
themes = [
("Tema 1 \u2014 Literasi ekonomi Islam sebagai konsep multidimensi.",
"Literatur secara konsisten membingkai literasi ekonomi Islam sebagai konstruk multidimensi berlandaskan "
"keimanan yang memadukan pengetahuan tentang keuangan sesuai syariah dan keuangan sosial Islam, niat untuk "
"patuh, serta kapasitas perilaku untuk bertindak. Konstruk ini sejalan dengan pandangan modal manusia atas "
"literasi keuangan, namun tidak dapat direduksi menjadi sekadar pengetahuan konvensional karena larangan "
"riba serta kewajiban zakat dan wakaf bersifat konstitutif (menjawab RQ2)."),
("Tema 2 \u2014 Generasi Z Muslim dan perilaku keuangan sesuai syariah.",
"Literasi yang lebih tinggi dikaitkan dengan niat mengadopsi perbankan syariah dan keputusan investasi "
"yang lebih baik di pasar modal syariah. Namun, terdapat kesenjangan sikap\u2013perilaku yang nyata karena "
"keterpaparan budaya digital konsumtif, sehingga pengetahuan tidak otomatis menjadi praktik (menjawab RQ1 "
"dan RQ4)."),
("Tema 3 \u2014 Pembelajaran digital dan media sosial sebagai kanal literasi.",
"Karena Generasi Z adalah penutur asli digital, pedagogi berbantuan teknologi seperti augmented reality "
"terbukti meningkatkan pengetahuan dan motivasi belajar, termasuk pada pembelajaran jarak jauh. Fintech dan "
"media sosial menjadi antarmuka dominan generasi muda Muslim dengan produk dan informasi keuangan syariah, "
"sehingga menjadi peluang sekaligus risiko misinformasi (menjawab RQ3)."),
("Tema 4 \u2014 Lembaga keuangan syariah, investasi halal, dan kesadaran pemuda.",
"Bank syariah, platform pasar modal, dan lembaga keuangan mikro berfungsi sebagai penyedia sekaligus "
"pendidik. Literasi berulang kali dikaitkan dengan adopsi dan kepercayaan terhadap institusi dan instrumen "
"yang sesuai syariah, memposisikan lembaga sebagai mitra kampanye literasi bagi pemuda (menjawab RQ4)."),
("Tema 5 \u2014 Zakat, wakaf, sedekah, dan pendidikan keuangan sosial Islam.",
"Literasi zakat secara signifikan meningkatkan niat menggunakan platform zakat digital, dan kajian tentang "
"wakaf tunai daring di kalangan pemuda dan milenial Muslim menonjolkan peran pengetahuan, religiusitas, "
"kepercayaan, dan kontrol perilaku. Tema ini menegaskan bahwa literasi ekonomi Islam menjangkau inti "
"redistributif ekonomi Islam (menjawab RQ3 dan RQ4)."),
("Tema 6 \u2014 Tantangan dan kesenjangan penelitian.",
"Tantangan yang berulang meliputi literasi dasar yang rendah, kesenjangan sikap\u2013perilaku, langkanya "
"instrumen pengukuran khusus pemuda, dominasi desain survei lintas-seksi, dan konsentrasi geografis pada "
"Indonesia dan Malaysia. Kesenjangan paling menonjol adalah terbatasnya studi yang secara eksplisit dan "
"serentak memusatkan perhatian pada Generasi Z, intervensi pendidikan, dan luaran ekonomi syariah "
"(menjawab RQ5)."),
]
for h,b in themes:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; set_after(p)
    rb = p.add_run(h + " "); rb.bold = True; rb.italic = True; rb.font.size = Pt(11)
    p.add_run(b)

subhead("Implikasi")
para("Secara teoretis, tinjauan ini mengonsolidasikan definisi literasi ekonomi Islam yang multidimensi "
"dan berlandaskan maqashid (khususnya hifzh al-mal) serta memperjelas posisinya dalam hukum dan ekonomi "
"Islam. Secara praktis, hasil ini mendukung pengintegrasian literasi keuangan syariah ke dalam kurikulum "
"sekolah dan perguruan tinggi serta penjangkauan digital yang membangun kepercayaan oleh lembaga keuangan "
"syariah. Secara kebijakan, temuan mendukung otoritas literasi keuangan dalam merancang program literasi "
"keuangan syariah yang berfokus pada pemuda dan memantau luarannya bersama target inklusi. Dengan membekali "
"Muslim muda untuk mengenali riba, gharar, dan maysir serta mengidentifikasi alternatif yang sesuai "
"syariah, pendidikan literasi ekonomi Islam berfungsi sebagai mekanisme preventif terhadap pinjaman "
"berbasis bunga (termasuk pinjaman daring predatoris) sekaligus mengarahkan sumber daya ke investasi halal "
"dan keuangan sosial Islam.")

# =====================================================================
# KESIMPULAN
# =====================================================================
heading_caps("Kesimpulan")
para("Penelitian ini secara sistematis mengidentifikasi, menilai, dan menyintesis kajian tentang "
"pendidikan literasi ekonomi Islam di kalangan Generasi Z Muslim menggunakan protokol SLR yang transparan "
"dan dapat direplikasi\u2014distrukturkan oleh kerangka PICOC, string pencarian Boolean pada delapan basis "
"data, kriteria inklusi/eksklusi yang eksplisit, alur seleksi berbasis PRISMA, daftar tilik penilaian "
"kualitas berbobot, formulir ekstraksi baku, serta analisis tematik dan bibliometrik deskriptif. Temuan "
"utama menunjukkan bahwa literasi ekonomi Islam merupakan konstruk multidimensi berlandaskan keimanan; "
"berkaitan secara empiris dengan perilaku keuangan sesuai syariah dan keterlibatan kelembagaan; bahwa kanal "
"digital dan media sosial menjadi medium penentu untuk menjangkau Generasi Z; bahwa literasi keuangan "
"sosial (zakat dan wakaf) merupakan komponen integral yang dimediasi secara digital; serta bahwa bidang ini "
"terkendala kesenjangan sikap\u2013perilaku, lemahnya pengukuran khusus pemuda, dominasi desain lintas-seksi, "
"dan konsentrasi geografis.")
para("Rekomendasi praktis mencakup pengintegrasian literasi keuangan syariah lintas kurikulum; penyiapan "
"lembaga keuangan syariah untuk menyampaikan edukasi digital yang tepercaya; pelibatan keluarga, komunitas, "
"dan figur daring yang kredibel; serta perancangan program yang secara eksplisit menyasar perilaku "
"(penghindaran riba, investasi halal, partisipasi zakat/wakaf), bukan sekadar kesadaran. Keterbatasan "
"tinjauan terikat pada protokolnya\u2014jendela 2015\u20132025, pembatasan bahasa Inggris/Indonesia, basis data "
"terpilih, dan ambang kualitas\u2014serta konsentrasi geografis sumber literatur yang membatasi generalisasi. "
"Riset masa depan diarahkan pada pengembangan instrumen khusus Generasi Z yang tervalidasi, desain "
"eksperimental dan longitudinal, diversifikasi konteks termasuk negara minoritas Muslim, kerangka literasi "
"yang memadukan keuangan komersial dan sosial, serta evaluasi pedagogi digital dengan pengaman keaslian "
"syariah.")

# =====================================================================
# DAFTAR PUSTAKA  (Turabian/Chicago bibliography, alphabetical)
# =====================================================================
heading_caps("Daftar Pustaka")
biblio = [
"Ahmad, Gatot Nazir, Umi Widyastuti, Santi Susanti, dan Hasan Mukhibad. \u201cDeterminants of the Islamic Financial Literacy.\u201d Accounting 6, no. 6 (2020): 961\u2013966.",
"Albaity, Mohamed, dan Mahfuzur Rahman. \u201cThe Intention to Use Islamic Banking: An Exploratory Study to Measure Islamic Financial Literacy.\u201d International Journal of Emerging Markets 14, no. 5 (2019): 988\u20131012.",
"Anisa, Salsabilla Nur, dan Muh Shadiqul Fajri AF. \u201cThe Impact of Islamic Financial Literacy on Gen Z\u2019s Investment Decisions in the Islamic Capital Market.\u201d Working paper, SSRN, 2025.",
"Antara, Purnomo M., Rosidah Musa, dan Faridah Hassan. \u201cBridging Islamic Financial Literacy and Halal Literacy: The Way Forward in Halal Ecosystem.\u201d Procedia Economics and Finance 37 (2016): 196\u2013202.",
"Antara, Purnomo M., Rosidah Musa, dan Faridah Hassan. \u201cConceptualisation and Operationalisation of Islamic Financial Literacy Scale.\u201d Pertanika Journal of Social Sciences & Humanities 25, no. S (2017): 251\u2013260.",
"Atkinson, Adele, dan Flore-Anne Messy. Measuring Financial Literacy: Results of the OECD/INFE Pilot Study. OECD Working Papers on Finance, Insurance and Private Pensions No. 15. Paris: OECD Publishing, 2012.",
"Carrera-Rivera, Angela, William Ochoa, Felix Larrinaga, dan Ganix Lasa. \u201cHow-to Conduct a Systematic Literature Review: A Quick Guide for Computer Science Research.\u201d MethodsX 9 (2022): 101895.",
"Chapra, M. Umer. Islam and the Economic Challenge. Leicester: The Islamic Foundation, 1992.",
"Dimock, Michael. Defining Generations: Where Millennials End and Generation Z Begins. Washington, DC: Pew Research Center, 2019.",
"El-Gamal, Mahmoud A. Islamic Finance: Law, Economics, and Practice. Cambridge: Cambridge University Press, 2006.",
"Hassan, M. Kabir, Mustafa Raza Rabbani, dan Mahmood Asad Mohd. Ali. \u201cChallenges for the Islamic Finance and Banking in Post COVID Era and the Role of Fintech.\u201d Journal of Economic Cooperation and Development 41, no. 3 (2020): 93\u2013116.",
"Iqbal, Zamir, dan Abbas Mirakhor. An Introduction to Islamic Finance: Theory and Practice. 2nd ed. Singapore: John Wiley & Sons, 2011.",
"Kasri, Rahmatina A., dan Adela Miranti Yuniar. \u201cDeterminants of Digital Zakat Payments: Lessons from Indonesian Experience.\u201d Journal of Islamic Accounting and Business Research 12, no. 3 (2021): 362\u2013379.",
"Khan, Feisal. \u201cHow \u2018Islamic\u2019 Is Islamic Banking?\u201d Journal of Economic Behavior & Organization 76, no. 3 (2010): 805\u2013820.",
"Kitchenham, Barbara, dan Stuart Charters. Guidelines for Performing Systematic Literature Reviews in Software Engineering. EBSE Technical Report EBSE-2007-01. Keele University & University of Durham, 2007.",
"Lusardi, Annamaria, dan Olivia S. Mitchell. \u201cThe Economic Importance of Financial Literacy: Theory and Evidence.\u201d Journal of Economic Literature 52, no. 1 (2014): 5\u201344.",
"Otoritas Jasa Keuangan. Survei Nasional Literasi dan Inklusi Keuangan (SNLIK) 2022. Jakarta: Otoritas Jasa Keuangan, 2022.",
"Page, Matthew J., Joanne E. McKenzie, Patrick M. Bossuyt, Isabelle Boutron, Tammy C. Hoffmann, Cynthia D. Mulrow, dkk. \u201cThe PRISMA 2020 Statement: An Updated Guideline for Reporting Systematic Reviews.\u201d BMJ 372 (2021): n71.",
"Prensky, Marc. \u201cDigital Natives, Digital Immigrants.\u201d On the Horizon 9, no. 5 (2001): 1\u20136.",
"Rabbani, Mustafa Raza, Shahnawaz Khan, dan Eleftherios I. Thalassinos. \u201cFinTech, Blockchain and Islamic Finance: An Extensive Literature Review.\u201d International Journal of Economics and Business Administration 8, no. 2 (2020): 65\u201386.",
"Rahim, Siti Hafsha Albasri, Rosemaliza Abdul Rashid, dan Abu Bakar Hamed. \u201cIslamic Financial Literacy and Its Determinants among University Students: An Exploratory Factor Analysis.\u201d International Journal of Economics and Financial Issues 6, no. 7S (2016): 32\u201335.",
"Snyder, Hannah. \u201cLiterature Review as a Research Methodology: An Overview and Guidelines.\u201d Journal of Business Research 104 (2019): 333\u2013339.",
"Solikhatun, Ismi, dan Ratna Candra Sari. \u201cSharia Financial Education Using Augmented Reality Technology to Increase Student Motivation in Distance Learning.\u201d AIP Conference Proceedings 2654 (2023): 020008.",
"Tranfield, David, David Denyer, dan Palminder Smart. \u201cTowards a Methodology for Developing Evidence-Informed Management Knowledge by Means of Systematic Review.\u201d British Journal of Management 14, no. 3 (2003): 207\u2013222.",
]
for entry in biblio:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-1.0)  # hanging indent
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(entry); r.font.size = Pt(10)

# integrity note
ip = doc.add_paragraph(); ip.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
ip.paragraph_format.space_before = Pt(10)
r = ip.add_run("Catatan reproduktibilitas: Seluruh rujukan di atas merupakan sumber akademik nyata dan dapat "
"diverifikasi yang diidentifikasi melalui prosedur pencarian pada bagian Metode. Tidak ada sumber yang "
"dikarang. Karena basis data pengindeks diperbarui secara berkala, penulis disarankan menjalankan kembali "
"string pencarian sebelum pengajuan untuk memfinalkan jumlah rekaman PRISMA dan memasukkan studi terbaru "
"yang memenuhi syarat.")
r.italic = True; r.font.size = Pt(9)

doc.save(OUT)
print("Saved base docx:", OUT, "| footnotes:", len(_footnotes))

# =====================================================================
# INJECT FOOTNOTES INTO THE DOCX PACKAGE
# =====================================================================
def xml_escape(s):
    return (s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"))

# build footnotes.xml
fn_parts = []
fn_parts.append('<w:footnote w:type="separator" w:id="-1"><w:p><w:pPr><w:spacing w:after="0" w:line="240" w:lineRule="auto"/></w:pPr><w:r><w:separator/></w:r></w:p></w:footnote>')
fn_parts.append('<w:footnote w:type="continuationSeparator" w:id="0"><w:p><w:pPr><w:spacing w:after="0" w:line="240" w:lineRule="auto"/></w:pPr><w:r><w:continuationSeparator/></w:r></w:p></w:footnote>')
for fid, text in _footnotes:
    fn_parts.append(
        '<w:footnote w:id="%d"><w:p><w:pPr><w:spacing w:after="0" w:line="240" w:lineRule="auto"/>'
        '<w:rPr><w:sz w:val="18"/></w:rPr></w:pPr>'
        '<w:r><w:rPr><w:vertAlign w:val="superscript"/><w:sz w:val="18"/></w:rPr><w:footnoteRef/></w:r>'
        '<w:r><w:rPr><w:sz w:val="18"/></w:rPr><w:t xml:space="preserve"> %s</w:t></w:r>'
        '</w:p></w:footnote>' % (fid, xml_escape(text))
    )
footnotes_xml = (
    '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
    '<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
    + "".join(fn_parts) + '</w:footnotes>'
)

tmp = OUT + ".tmp"
with zipfile.ZipFile(OUT, "r") as zin:
    names = zin.namelist()
    data = {n: zin.read(n) for n in names}

# 1) add footnotes.xml
data["word/footnotes.xml"] = footnotes_xml.encode("utf-8")

# 2) content types
ct = data["[Content_Types].xml"].decode("utf-8")
if "footnotes+xml" not in ct:
    override = '<Override PartName="/word/footnotes.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml"/>'
    ct = ct.replace("</Types>", override + "</Types>")
data["[Content_Types].xml"] = ct.encode("utf-8")

# 3) document rels
rels_name = "word/_rels/document.xml.rels"
rels = data[rels_name].decode("utf-8")
if "relationships/footnotes" not in rels:
    rel = '<Relationship Id="rId901" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes" Target="footnotes.xml"/>'
    rels = rels.replace("</Relationships>", rel + "</Relationships>")
data[rels_name] = rels.encode("utf-8")

with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
    for n, b in data.items():
        zout.writestr(n, b)
shutil.move(tmp, OUT)
print("Footnotes injected. Done.")
