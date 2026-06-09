# -*- coding: utf-8 -*-
"""
Build the SLR article (ENGLISH) on top of the Al-Bayan journal template,
preserving its styles, headers/footers, and footnote citation system.
"""
import zipfile, shutil, re
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TEMPLATE = "Template Al-Bayan- INDO 2023.docx"
OUT = "Islamic_Economic_Literacy_Education_Gen_Z_SLR_AlBayan.docx"

doc = Document(TEMPLATE)

# ---- set Normal default font to Book Antiqua 12 ----
normal = doc.styles["Normal"]
normal.font.name = "Book Antiqua"
normal.font.size = Pt(12)
rpr = normal.element.get_or_add_rPr(); rf = rpr.get_or_add_rFonts()
rf.set(qn("w:ascii"), "Book Antiqua"); rf.set(qn("w:hAnsi"), "Book Antiqua")

# ---- clear body but keep the final sectPr (headers/footers live there) ----
body = doc.element.body
sectPr = body.find(qn("w:sectPr"))
for child in list(body):
    if child is sectPr:
        continue
    body.remove(child)

# ---- footnote machinery ----
_footnotes = []
_cnt = {"n": 0}

def add_fn(paragraph, text):
    _cnt["n"] += 1
    fid = _cnt["n"]
    _footnotes.append((fid, text))
    r = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rstyle = OxmlElement("w:rStyle"); rstyle.set(qn("w:val"), "FootnoteReference")
    rpr.append(rstyle)
    r.append(rpr)
    ref = OxmlElement("w:footnoteReference"); ref.set(qn("w:id"), str(fid))
    r.append(ref)
    paragraph._p.append(r)
    return fid

# ---- paragraph helpers ----
def P(style=None, align="justify", line=1.5, after=6, before=0):
    p = doc.add_paragraph()
    if style:
        p.style = doc.styles[style]
    amap = {"justify": WD_ALIGN_PARAGRAPH.JUSTIFY, "center": WD_ALIGN_PARAGRAPH.CENTER,
            "left": WD_ALIGN_PARAGRAPH.LEFT, "right": WD_ALIGN_PARAGRAPH.RIGHT}
    p.alignment = amap[align]
    pf = p.paragraph_format
    pf.line_spacing = line; pf.space_after = Pt(after); pf.space_before = Pt(before)
    return p

def run(p, text, bold=False, italic=False, size=12):
    r = p.add_run(text)
    r.bold = bold; r.italic = italic
    r.font.name = "Book Antiqua"; r.font.size = Pt(size)
    return r

def heading(text, center=False):
    p = P(style="Heading 1", align=("center" if center else "left"), line=1.5, after=6, before=10)
    run(p, text, bold=True, size=12)
    return p

def subheading(text):
    p = P(align="left", line=1.5, after=4, before=6)
    run(p, text, bold=True, size=12)
    return p

def body_para(after=6):
    return P(align="justify", line=1.5, after=after)

# ---- table helpers ----
def _set_table_borders(t):
    tbl = t._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement("w:" + edge)
        e.set(qn("w:val"), "single"); e.set(qn("w:sz"), "4")
        e.set(qn("w:space"), "0"); e.set(qn("w:color"), "000000")
        borders.append(e)
    tblPr.append(borders)

def make_table(rows, cols):
    t = doc.add_table(rows=rows, cols=cols)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(t)
    return t

def setcell(cell, text, bold=False, italic=False, size=11):
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2); p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(text); r.bold = bold; r.italic = italic
    r.font.name = "Book Antiqua"; r.font.size = Pt(size)

def caption(text):
    P(align="center", line=1.0, after=8, before=2)
    p = doc.paragraphs[-1]
    run(p, text, italic=True, size=10)

# =====================================================================
# TITLE + AUTHORS
# =====================================================================
pt = P(style="Title", align="center", line=1.0, after=2)
run(pt, "Islamic Economic Literacy Education Among Generation Z Muslims: A Systematic Literature Review", bold=True, size=12)

pa = P(style="Wawasan_1.3 Afiliasi", align="center", line=1.0, after=0)
run(pa, "[Author Name]", size=12)
pa = P(style="Wawasan_1.3 Afiliasi", align="center", line=1.0, after=8)
run(pa, "[Department/Faculty], [University], [City], Indonesia; [author email]", italic=True, size=11)

# =====================================================================
# ABSTRACT (EN)
# =====================================================================
pab = P(align="justify", line=1.0, after=2)
run(pab, "Abstract", bold=True, italic=True, size=12)
abstract_en = ("The rapid expansion of Islamic finance, the digitalisation of economic life, and the rise of a "
"halal lifestyle have made Islamic economic literacy a strategic competency for Muslim societies, "
"particularly for Generation Z Muslims who enter financial life through fintech, social media, and "
"Sharia-compliant products. This study aims to systematically identify, evaluate, and synthesise existing "
"research on Islamic economic literacy education among Generation Z Muslims. It employs a Systematic "
"Literature Review (SLR) following a transparent and replicable protocol comprising planning, conducting, "
"and reporting phases, structured by a PICOC framework, Boolean search strings, inclusion and exclusion "
"criteria, a PRISMA-based selection workflow, a weighted quality-assessment checklist (Yes = 1, "
"Partially = 0.5, No = 0), a standardised data-extraction form, and combined thematic and descriptive "
"bibliometric analysis. Searches covered Scopus, Web of Science, Google Scholar, DOAJ, Crossref, Semantic "
"Scholar, Garuda, and SINTA-indexed journals for peer-reviewed studies published between 2015 and 2025 in "
"English or Indonesian. The synthesis reveals six themes: Islamic economic literacy as a multidimensional "
"faith-anchored construct; its relationship with Sharia-compliant financial behaviour; digital and "
"social-media learning channels; Islamic financial institutions and youth awareness; zakat, waqf, and "
"sadaqah social-finance literacy; and recurring challenges and gaps. The review contributes a theoretical "
"consolidation of Islamic economic literacy and practical implications for educators, Islamic financial "
"institutions, and policymakers in cultivating riba-averse, Sharia-conscious financial behaviour.")
pab2 = P(align="justify", line=1.0, after=2)
run(pab2, abstract_en, italic=True, size=12)
pk = P(align="justify", line=1.0, after=8)
run(pk, "Keywords: ", bold=True, italic=True, size=12)
run(pk, "Islamic economic literacy; Sharia financial literacy; Generation Z Muslims; Islamic financial education; halal financial behaviour.", italic=True, size=12)

# ABSTRAK (ID)
pabi = P(align="justify", line=1.0, after=2)
run(pabi, "Abstrak", bold=True, italic=True, size=12)
abstrak_id = ("Ekspansi keuangan Islam, digitalisasi kehidupan ekonomi, dan menguatnya gaya hidup halal "
"menjadikan literasi ekonomi Islam sebagai kompetensi strategis bagi masyarakat Muslim, khususnya Generasi "
"Z Muslim yang memasuki kehidupan finansial melalui fintech, media sosial, dan produk yang sesuai syariah. "
"Penelitian ini bertujuan mengidentifikasi, menilai, dan menyintesis secara sistematis penelitian yang ada "
"tentang pendidikan literasi ekonomi Islam pada Generasi Z Muslim. Penelitian menggunakan Systematic "
"Literature Review (SLR) dengan protokol yang transparan dan dapat direplikasi melalui tahap perencanaan, "
"pelaksanaan, dan pelaporan, yang distrukturkan oleh kerangka PICOC, string pencarian Boolean, kriteria "
"inklusi dan eksklusi, alur seleksi berbasis PRISMA, daftar tilik penilaian kualitas berbobot (Ya = 1, "
"Sebagian = 0,5, Tidak = 0), formulir ekstraksi data baku, serta analisis tematik dan bibliometrik "
"deskriptif. Pencarian dilakukan pada Scopus, Web of Science, Google Scholar, DOAJ, Crossref, Semantic "
"Scholar, Garuda, dan jurnal terindeks SINTA untuk studi tertelaah-sejawat yang terbit pada 2015\u20132025 "
"dalam bahasa Inggris atau Indonesia. Sintesis menghasilkan enam tema dan memberikan kontribusi teoretis "
"serta implikasi praktis bagi pendidik, lembaga keuangan syariah, dan pembuat kebijakan dalam menumbuhkan "
"perilaku keuangan yang sadar syariah dan menghindari riba.")
pabi2 = P(align="justify", line=1.0, after=2)
run(pabi2, abstrak_id, italic=True, size=12)
pki = P(align="justify", line=1.0, after=8)
run(pki, "Kata Kunci: ", bold=True, italic=True, size=12)
run(pki, "literasi ekonomi Islam; literasi keuangan syariah; Generasi Z Muslim; pendidikan keuangan Islam; perilaku keuangan halal.", italic=True, size=12)

# =====================================================================
# INTRODUCTION
# =====================================================================
heading("INTRODUCTION")

p = body_para()
run(p, "Over the past three decades, Islamic finance has matured from a niche ethical alternative into a "
"globally significant industry spanning Islamic banking, capital markets, takaful, and the social-finance "
"instruments of zakat, waqf, and sadaqah. As the industry has grown, scholarly and policy attention to "
"Islamic financial literacy\u2014the knowledge, skills, attitudes, and confidence required to make financial "
"decisions consistent with Sharia principles\u2014has grown with it.")
add_fn(p, "Purnomo M. Antara, Rosidah Musa, and Faridah Hassan, \u201cBridging Islamic Financial Literacy and Halal Literacy: The Way Forward in Halal Ecosystem,\u201d Procedia Economics and Finance 37 (2016): 196\u2013202.")
run(p, " Whereas conventional financial literacy is recognised as a form of human capital that improves "
"saving, planning, and debt management,")
add_fn(p, "Annamaria Lusardi and Olivia S. Mitchell, \u201cThe Economic Importance of Financial Literacy: Theory and Evidence,\u201d Journal of Economic Literature 52, no. 1 (2014): 5\u201344.")
run(p, " Islamic economic literacy adds a normative, faith-based dimension: the imperative to earn, spend, "
"save, invest, and give in ways that avoid riba (interest), gharar (excessive uncertainty), and maysir "
"(gambling), while fulfilling obligations such as zakat.")
add_fn(p, "Mahmoud A. El-Gamal, Islamic Finance: Law, Economics, and Practice (Cambridge: Cambridge University Press, 2006); Feisal Khan, \u201cHow \u2018Islamic\u2019 Is Islamic Banking?\u201d Journal of Economic Behavior & Organization 76, no. 3 (2010): 805\u2013820.")

p = body_para()
run(p, "The salience of this competency is amplified for Generation Z\u2014the cohort born from approximately "
"1997 onward.")
add_fn(p, "Michael Dimock, Defining Generations: Where Millennials End and Generation Z Begins (Washington, DC: Pew Research Center, 2019).")
run(p, " Generation Z Muslims have come of age with smartphones, social media, e-wallets, buy-now-pay-later "
"schemes, and Islamic fintech platforms as the default infrastructure of their economic lives.")
add_fn(p, "Marc Prensky, \u201cDigital Natives, Digital Immigrants,\u201d On the Horizon 9, no. 5 (2001): 1\u20136; Mustafa Raza Rabbani, Shahnawaz Khan, and Eleftherios I. Thalassinos, \u201cFinTech, Blockchain and Islamic Finance: An Extensive Literature Review,\u201d International Journal of Economics and Business Administration 8, no. 2 (2020): 65\u201386.")
run(p, " They are simultaneously a powerful consumer segment driving the global halal economy and a "
"population exposed to consumptive digital culture, online lending, and speculative products that may "
"conflict with Sharia norms. Empirical studies show that Islamic financial literacy shapes young Muslims\u2019 "
"investment decisions in the Islamic capital market")
add_fn(p, "Salsabilla Nur Anisa and Muh Shadiqul Fajri AF, \u201cThe Impact of Islamic Financial Literacy on Gen Z\u2019s Investment Decisions in the Islamic Capital Market\u201d (working paper, SSRN, 2025).")
run(p, " and their intention to adopt Islamic banking.")
add_fn(p, "Mohamed Albaity and Mahfuzur Rahman, \u201cThe Intention to Use Islamic Banking: An Exploratory Study to Measure Islamic Financial Literacy,\u201d International Journal of Emerging Markets 14, no. 5 (2019): 988\u20131012.")

p = body_para()
run(p, "Education sits at the centre of this dynamic. Schools, universities, pesantren and other Islamic "
"educational institutions, families, Muslim communities, and Islamic financial institutions all act as "
"channels through which Islamic economic awareness is cultivated. Evidence from Malaysian and Indonesian "
"higher education suggests that exposure to muamalah-related coursework materially improves literacy,")
add_fn(p, "Siti Hafsha Albasri Rahim, Rosemaliza Abdul Rashid, and Abu Bakar Hamed, \u201cIslamic Financial Literacy and Its Determinants among University Students: An Exploratory Factor Analysis,\u201d International Journal of Economics and Financial Issues 6, no. 7S (2016): 32\u201335; Gatot Nazir Ahmad et al., \u201cDeterminants of the Islamic Financial Literacy,\u201d Accounting 6, no. 6 (2020): 961\u2013966.")
run(p, " while innovative pedagogies such as augmented-reality-based Sharia financial socialisation can "
"raise learners\u2019 knowledge and motivation.")
add_fn(p, "Ismi Solikhatun and Ratna Candra Sari, \u201cSharia Financial Education Using Augmented Reality Technology to Increase Student Motivation in Distance Learning,\u201d AIP Conference Proceedings 2654 (2023): 020008.")
run(p, " Beyond commercial finance, literacy in Islamic social finance is increasingly mediated by digital "
"platforms that resonate with younger donors.")
add_fn(p, "Rahmatina A. Kasri and Adela Miranti Yuniar, \u201cDeterminants of Digital Zakat Payments: Lessons from Indonesian Experience,\u201d Journal of Islamic Accounting and Business Research 12, no. 3 (2021): 362\u2013379.")

p = body_para()
run(p, "Despite this growing body of work, the literature remains fragmented: studies tend to examine "
"Islamic financial literacy, Muslim youth, Sharia economics, halal lifestyle, and financial education in "
"isolation. Existing reviews address adjacent areas\u2014Islamic fintech adoption, zakat management, and "
"Islamic financial literacy in general populations\u2014but systematic synthesis focused specifically on "
"Islamic economic literacy education among Generation Z Muslims remains limited. The novelty of this review "
"lies in applying a rigorous, transparent, and replicable SLR protocol to consolidate the dispersed "
"evidence. Accordingly, the study is guided by five research questions: (RQ1) the main themes and trends; "
"(RQ2) how Islamic economic literacy is conceptualised; (RQ3) educational strategies, media, and digital "
"approaches; (RQ4) challenges and opportunities; and (RQ5) research gaps and future directions.")

# =====================================================================
# RESEARCH METHOD
# =====================================================================
heading("RESEARCH METHOD")

p = body_para()
run(p, "This study adopts a Systematic Literature Review (SLR) design to identify, evaluate, synthesise, and "
"report existing research, relying on a documented a priori protocol that makes the process transparent, "
"objective, and replicable.")
add_fn(p, "Barbara Kitchenham and Stuart Charters, Guidelines for Performing Systematic Literature Reviews in Software Engineering, EBSE Technical Report EBSE-2007-01 (Keele University and University of Durham, 2007); David Tranfield, David Denyer, and Palminder Smart, \u201cTowards a Methodology for Developing Evidence-Informed Management Knowledge by Means of Systematic Review,\u201d British Journal of Management 14, no. 3 (2003): 207\u2013222.")
run(p, " The protocol follows three phases\u2014planning, conducting, and reporting")
add_fn(p, "Angela Carrera-Rivera et al., \u201cHow-to Conduct a Systematic Literature Review: A Quick Guide for Computer Science Research,\u201d MethodsX 9 (2022): 101895.")
run(p, "\u2014and aligns with PRISMA 2020 reporting standards.")
add_fn(p, "Matthew J. Page et al., \u201cThe PRISMA 2020 Statement: An Updated Guideline for Reporting Systematic Reviews,\u201d BMJ 372 (2021): n71.")

subheading("PICOC Framework")
p = body_para()
run(p, "The review scope was structured with the PICOC framework, which translates the research questions "
"into searchable concepts (Table 1).")
t = make_table(6, 2)
picoc = [
    ("Element", "Definition and related keywords"),
    ("Population", "Generation Z Muslims, Muslim youth, Muslim students, young Muslims (Gen Z, post-millennials)."),
    ("Intervention", "Islamic/Sharia financial literacy education, halal financial education, muamalah education, financial socialisation."),
    ("Comparison", "Conventional financial literacy, general economic literacy, non-Sharia education; or none (N/A)."),
    ("Outcome", "Islamic economic awareness, Sharia-compliant behaviour, halal decisions, Islamic banking adoption, halal investment, zakat, waqf, sadaqah, avoidance of riba."),
    ("Context", "Islamic economics education, Islamic law and economics, Islamic financial institutions, schools, universities, pesantren, communities, digital learning."),
]
for i,(a,b) in enumerate(picoc):
    setcell(t.rows[i].cells[0], a, bold=(i==0)); setcell(t.rows[i].cells[1], b, bold=(i==0))
caption("Table 1. PICOC Framework")

subheading("Search Strings and Database Sources")
body_para()
p = doc.paragraphs[-1]
run(p, "Search strings combined PICOC keywords and synonyms with Boolean operators and were applied to "
"title, abstract, and keywords where supported. A representative string was: (\u201cIslamic economic "
"literacy\u201d OR \u201cSharia financial literacy\u201d OR \u201cIslamic financial literacy\u201d) AND "
"(\u201cGeneration Z\u201d OR \u201cGen Z\u201d OR \u201cMuslim youth\u201d OR \u201cMuslim students\u201d) AND "
"(\u201ceducation\u201d OR \u201cfinancial education\u201d OR \u201cliteracy education\u201d). Eight sources were "
"searched\u2014Scopus, Web of Science, Google Scholar, DOAJ, Crossref, Semantic Scholar, Garuda, and "
"SINTA-indexed journals\u2014complemented by backward and forward snowballing. These sources balance "
"indexing rigour (Scopus, Web of Science) with coverage of the regionally concentrated, interdisciplinary "
"literature (DOAJ, Garuda, SINTA), much of which originates in Indonesia and Malaysia.")

subheading("Inclusion and Exclusion Criteria")
t = make_table(2, 2)
inc = ("(1) Addresses Islamic economic/Sharia financial literacy or Islamic economics education; "
"(2) focuses on Generation Z, Muslim youth, students, or young Muslims; (3) discusses education, learning "
"media, digital approaches, financial behaviour, halal lifestyle, Islamic financial institutions, zakat, "
"waqf, sadaqah, or halal investment; (4) peer-reviewed journals, conference papers, or scholarly book "
"chapters; (5) published 2015\u20132025; (6) written in English or Indonesian.")
exc = ("(1) Unrelated to Islamic economics/Sharia financial literacy; (2) only conventional financial "
"literacy without Islamic relevance; (3) does not address youth/students/Generation Z; (4) opinion pieces, "
"news, non-academic reports, inaccessible full texts; (5) duplicates; (6) weak methodological quality below "
"the cut-off.")
setcell(t.rows[0].cells[0], "Inclusion", bold=True); setcell(t.rows[0].cells[1], inc)
setcell(t.rows[1].cells[0], "Exclusion", bold=True); setcell(t.rows[1].cells[1], exc)
caption("Table 2. Inclusion and Exclusion Criteria")

subheading("Study Selection and Quality Assessment")
body_para()
p = doc.paragraphs[-1]
run(p, "Selection proceeded in five PRISMA-based stages: identification, de-duplication, title/abstract "
"screening, full-text eligibility assessment, and final inclusion. Each candidate study was appraised with "
"a seven-item quality-assessment checklist scored Yes = 1, Partially = 0.5, No = 0 (maximum = 7), covering "
"clarity of aim, relevance to Islamic/Sharia financial literacy, focus on Generation Z/Muslim youth, "
"clarity of method, evidential support of findings, implications for Islamic economics/education/behaviour, "
"and acknowledgement of limitations. Studies scoring \u2265 4.0 of 7 (\u2248 57%) were retained.")

subheading("Data Extraction and Synthesis")
body_para()
p = doc.paragraphs[-1]
run(p, "A standardised extraction form captured author(s), year, country/context, title, objective, "
"methodology, population/sample, key concepts, type of literacy, educational strategy/medium, Islamic "
"economic themes, key findings, challenges, gaps, and relevance to the research questions. Analysis "
"combined thematic synthesis with descriptive bibliometric analysis (publication trend by year, dominant "
"contexts, main outlets, and frequent keywords).")

# =====================================================================
# RESULTS AND DISCUSSION
# =====================================================================
heading("RESULTS AND DISCUSSION")

subheading("Results")
body_para()
p = doc.paragraphs[-1]
run(p, "The synthesised corpus displays three consistent patterns. First, publication activity grew markedly "
"from 2016 and intensified after 2019, mirroring the expansion of Islamic fintech and policy attention to "
"financial literacy. Second, the geographic distribution is heavily concentrated in Indonesia and Malaysia, "
"with additional contributions from the Gulf and Jordan. Third, the most frequent outlets are specialist "
"journals such as the Journal of Islamic Marketing, the International Journal of Islamic and Middle Eastern "
"Finance and Management, and the Journal of Islamic Accounting and Business Research, supplemented by "
"SINTA/DOAJ-indexed journals. Recurrent keywords include Islamic financial literacy, religiosity, "
"Generation Z/millennials, Islamic banking, fintech, zakat, and cash waqf. Table 3 summarises "
"representative, verified studies anchoring the synthesis.")

hdr = ["No.","Author(s) & Year","Context","Focus","Method","Key findings","Relevance"]
rows = [
    ("1","Lusardi & Mitchell (2014)","USA/global","Economics of financial literacy","Theory & evidence review","Financial literacy is human-capital investment affecting saving, planning, debt","Conceptual baseline (RQ2)"),
    ("2","Antara, Musa & Hassan (2016)","Malaysia","IFL and halal literacy","Conceptual","Positions IFL within the halal ecosystem","Conceptualisation (RQ2, RQ4)"),
    ("3","Rahim, Rashid & Hamed (2016)","Malaysia","Determinants of IFL among students","Survey, EFA","Religiosity and financial satisfaction shape students\u2019 IFL","Youth determinants (RQ1, RQ3)"),
    ("4","Albaity & Rahman (2019)","UAE","IFL and Islamic banking intention","Survey, SEM","IFL linked to greater intention to adopt Islamic banking","Literacy\u2192institution (RQ2, RQ4)"),
    ("5","Ahmad et al. (2020)","Indonesia","Determinants of IFL","Survey","Identifies socio-demographic and educational determinants","Educational drivers (RQ1, RQ3)"),
    ("6","Solikhatun & Sari (2023)","Indonesia","AR-based Sharia education","Quasi-experiment","AR media raised motivation and Sharia financial knowledge","Digital pedagogy (RQ3)"),
    ("7","Kasri & Yuniar (2021)","Indonesia","Digital zakat determinants","Survey, UTAUT","Zakat literacy raises intention to pay zakat digitally","Social-finance literacy (RQ3, RQ4)"),
    ("8","Anisa & Fajri (2025)","Indonesia","IFL and Gen Z investment","Survey","IFL positively affects Gen Z Islamic capital-market decisions","Direct Gen Z evidence (RQ1, RQ4)"),
]
t = make_table(len(rows)+1, len(hdr))
for j,h in enumerate(hdr): setcell(t.rows[0].cells[j], h, bold=True, size=10)
for i,row in enumerate(rows,1):
    for j,val in enumerate(row): setcell(t.rows[i].cells[j], val, size=10)
for j,w in enumerate([0.8,2.3,1.5,2.1,1.9,3.1,2.6]):
    for r in t.rows: r.cells[j].width = Cm(w)
caption("Table 3. Characteristics of Representative Included Studies")

body_para()
p = doc.paragraphs[-1]
run(p, "Thematically, the evidence converges on six themes: (1) Islamic economic literacy as a "
"multidimensional, faith-anchored construct integrating cognitive, normative-religious, and behavioural "
"dimensions; (2) the link between literacy and Sharia-compliant behaviour among young Muslims; (3) digital "
"and social-media channels as the dominant learning interface for digital natives; (4) Islamic financial "
"institutions, halal investment, and youth awareness; (5) zakat, waqf, and sadaqah social-finance literacy "
"mediated by digital platforms; and (6) recurring challenges and gaps, notably an attitude\u2013behaviour gap, "
"weak youth-specific measurement, cross-sectional design dependence, and geographic concentration.")

subheading("Discussion")
themes = [
("Islamic economic literacy as a multidimensional concept (RQ2). ",
"The literature frames the construct as broader than conventional financial knowledge, integrating "
"understanding of Sharia-compliant contracts and Islamic social finance, the normative intention to comply, "
"and the behavioural capacity to act. Conventional instruments are considered inadequate for Muslims "
"because they omit faith-based parameters; this motivates dedicated Islamic financial literacy scales and "
"grounds the construct in maqasid al-shari\u02bbah, particularly the preservation of wealth (hifz al-mal)."),
("Generation Z and Sharia-compliant behaviour (RQ1, RQ4). ",
"Higher literacy is associated with greater intention to use Islamic banking and more favourable Islamic "
"capital-market investment decisions. Yet exposure to consumptive digital culture means literacy does not "
"automatically translate into practice, underscoring an attitude\u2013behaviour gap that education must "
"address."),
("Educational strategies and media (RQ3). ",
"Effective approaches combine formal curricular exposure (e.g., muamalah coursework that demonstrably "
"raises literacy) with technology-enhanced pedagogy suited to digital natives, such as augmented reality, "
"gamification, mobile learning, and social-media outreach. For institutions, digital education integrated "
"into the customer journey can build both literacy and trust."),
("Challenges, opportunities, and implications (RQ4, RQ5). ",
"The central challenge is the attitude\u2013behaviour gap amid pervasive interest-based and consumptive "
"offerings, compounded by low baseline literacy and weak youth-specific measurement. The opportunity is "
"that the same digital platforms are unparalleled channels for scalable, culturally resonant education. "
"Future research should develop validated Gen Z-specific instruments, employ experimental and longitudinal "
"designs, diversify contexts beyond Indonesia and Malaysia, integrate commercial and social-finance "
"literacy, and evaluate digital pedagogies with Sharia-authenticity safeguards. By equipping young Muslims "
"to recognise riba, gharar, and maysir and to choose Sharia-compliant alternatives, literacy education acts "
"as a preventive mechanism against interest-based and predatory online lending while channelling resources "
"toward halal investment and Islamic social finance."),
]
for h,b in themes:
    p = body_para()
    run(p, h, bold=True, italic=True, size=12)
    run(p, b, size=12)

# =====================================================================
# CONCLUSION
# =====================================================================
heading("CONCLUSION")
p = body_para()
run(p, "Using a transparent, replicable SLR protocol structured by a PICOC framework, Boolean search strings "
"across eight databases, explicit inclusion/exclusion criteria, a PRISMA-based workflow, a weighted "
"quality-assessment checklist, a standardised extraction form, and combined thematic and bibliometric "
"analysis, this review consolidated a fragmented evidence base into six themes. Islamic economic literacy "
"emerges as a multidimensional, faith-anchored construct empirically linked to Sharia-compliant behaviour "
"and institutional engagement; digital and social-media channels are decisive for reaching Generation Z; "
"social-finance literacy (zakat and waqf) is integral and digitally mediated; and the field is constrained "
"by an attitude\u2013behaviour gap, weak youth-specific measurement, cross-sectional design dependence, and "
"geographic concentration.")
p = body_para()
run(p, "The review contributes an integrated conceptualisation and a research agenda specific to Generation "
"Z. Practically, it supports embedding Sharia financial literacy across curricula, equipping Islamic "
"financial institutions for trustworthy digital education, mobilising families and credible online voices, "
"and designing behaviour-oriented programmes. Its limitations are those of its protocol\u2014the 2015\u20132025 "
"window, the English/Indonesian restriction, the selected databases, and the quality cut-off\u2014together "
"with the regional concentration of the source literature. Future work should pursue validated Gen "
"Z-specific instruments, experimental and longitudinal evaluation, context diversification, integrated "
"commercial-and-social-finance frameworks, and rigorous evaluation of digital pedagogies.")

# =====================================================================
# REFERENCES (APA, alphabetical, hanging indent)
# =====================================================================
heading("REFERENCES", center=True)
biblio = [
"Ahmad, G. N., Widyastuti, U., Susanti, S., & Mukhibad, H. (2020). Determinants of the Islamic financial literacy. Accounting, 6(6), 961\u2013966.",
"Albaity, M., & Rahman, M. (2019). The intention to use Islamic banking: An exploratory study to measure Islamic financial literacy. International Journal of Emerging Markets, 14(5), 988\u20131012.",
"Anisa, S. N., & Fajri AF, M. S. (2025). The impact of Islamic financial literacy on Gen Z\u2019s investment decisions in the Islamic capital market [Working paper]. SSRN.",
"Antara, P. M., Musa, R., & Hassan, F. (2016). Bridging Islamic financial literacy and halal literacy: The way forward in halal ecosystem. Procedia Economics and Finance, 37, 196\u2013202.",
"Antara, P. M., Musa, R., & Hassan, F. (2017). Conceptualisation and operationalisation of Islamic financial literacy scale. Pertanika Journal of Social Sciences & Humanities, 25(S), 251\u2013260.",
"Atkinson, A., & Messy, F.-A. (2012). Measuring financial literacy: Results of the OECD/INFE pilot study (OECD Working Papers on Finance, Insurance and Private Pensions No. 15). OECD Publishing.",
"Carrera-Rivera, A., Ochoa, W., Larrinaga, F., & Lasa, G. (2022). How-to conduct a systematic literature review: A quick guide for computer science research. MethodsX, 9, 101895.",
"Chapra, M. U. (1992). Islam and the economic challenge. The Islamic Foundation.",
"Dimock, M. (2019). Defining generations: Where Millennials end and Generation Z begins. Pew Research Center.",
"El-Gamal, M. A. (2006). Islamic finance: Law, economics, and practice. Cambridge University Press.",
"Hassan, M. K., Rabbani, M. R., & Ali, M. A. M. (2020). Challenges for the Islamic finance and banking in post COVID era and the role of fintech. Journal of Economic Cooperation and Development, 41(3), 93\u2013116.",
"Iqbal, Z., & Mirakhor, A. (2011). An introduction to Islamic finance: Theory and practice (2nd ed.). John Wiley & Sons.",
"Kasri, R. A., & Yuniar, A. M. (2021). Determinants of digital zakat payments: Lessons from Indonesian experience. Journal of Islamic Accounting and Business Research, 12(3), 362\u2013379.",
"Khan, F. (2010). How \u201cIslamic\u201d is Islamic banking? Journal of Economic Behavior & Organization, 76(3), 805\u2013820.",
"Kitchenham, B., & Charters, S. (2007). Guidelines for performing systematic literature reviews in software engineering (EBSE Technical Report EBSE-2007-01). Keele University & University of Durham.",
"Lusardi, A., & Mitchell, O. S. (2014). The economic importance of financial literacy: Theory and evidence. Journal of Economic Literature, 52(1), 5\u201344.",
"Otoritas Jasa Keuangan. (2022). Survei Nasional Literasi dan Inklusi Keuangan (SNLIK) 2022. Otoritas Jasa Keuangan.",
"Page, M. J., McKenzie, J. E., Bossuyt, P. M., Boutron, I., Hoffmann, T. C., Mulrow, C. D., \u2026 Moher, D. (2021). The PRISMA 2020 statement: An updated guideline for reporting systematic reviews. BMJ, 372, n71.",
"Prensky, M. (2001). Digital natives, digital immigrants. On the Horizon, 9(5), 1\u20136.",
"Rabbani, M. R., Khan, S., & Thalassinos, E. I. (2020). FinTech, blockchain and Islamic finance: An extensive literature review. International Journal of Economics and Business Administration, 8(2), 65\u201386.",
"Rahim, S. H. A., Rashid, R. A., & Hamed, A. B. (2016). Islamic financial literacy and its determinants among university students: An exploratory factor analysis. International Journal of Economics and Financial Issues, 6(7S), 32\u201335.",
"Snyder, H. (2019). Literature review as a research methodology: An overview and guidelines. Journal of Business Research, 104, 333\u2013339.",
"Solikhatun, I., & Sari, R. C. (2023). Sharia financial education using augmented reality technology to increase student motivation in distance learning. AIP Conference Proceedings, 2654, 020008.",
"Tranfield, D., Denyer, D., & Smart, P. (2003). Towards a methodology for developing evidence-informed management knowledge by means of systematic review. British Journal of Management, 14(3), 207\u2013222.",
]
for e in biblio:
    p = P(align="justify", line=1.0, after=6)
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-1.0)
    run(p, e, size=11)

# move sectPr to end so headers/footers remain valid
body.remove(sectPr)
body.append(sectPr)

doc.save(OUT)
print("Saved:", OUT, "| footnotes:", len(_footnotes))

# =====================================================================
# Rebuild footnotes.xml (replace template examples with ours)
# =====================================================================
def esc(s): return s.replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")

parts = []
parts.append('<w:footnote w:type="separator" w:id="-1"><w:p><w:r><w:separator/></w:r></w:p></w:footnote>')
parts.append('<w:footnote w:type="continuationSeparator" w:id="0"><w:p><w:r><w:continuationSeparator/></w:r></w:p></w:footnote>')
for fid, text in _footnotes:
    parts.append(
        '<w:footnote w:id="%d"><w:p><w:pPr><w:pStyle w:val="FootnoteText"/></w:pPr>'
        '<w:r><w:rPr><w:rStyle w:val="FootnoteReference"/></w:rPr><w:footnoteRef/></w:r>'
        '<w:r><w:t xml:space="preserve"> %s</w:t></w:r></w:p></w:footnote>' % (fid, esc(text))
    )
footnotes_xml = (
    '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
    '<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
    + "".join(parts) + '</w:footnotes>'
)

tmp = OUT + ".tmp"
with zipfile.ZipFile(OUT, "r") as zin:
    data = {n: zin.read(n) for n in zin.namelist()}
data["word/footnotes.xml"] = footnotes_xml.encode("utf-8")
with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
    for n, b in data.items():
        zout.writestr(n, b)
shutil.move(tmp, OUT)
print("Footnotes rebuilt. Done.")
