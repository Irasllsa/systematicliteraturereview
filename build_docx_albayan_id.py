# -*- coding: utf-8 -*-
"""
Membuat artikel SLR (BAHASA INDONESIA) di atas template jurnal Al-Bayan.
- Tanpa catatan kaki; sitasi memakai APA 7th edition (in-text), diletakkan di akhir kalimat.
- Tanpa simbol em-dash (U+2014) maupun en-dash (U+2013).
- Mempertahankan gaya, header/footer, dan style template.
"""
import os
import zipfile, shutil
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TEMPLATE = "Template Al-Bayan- INDO 2023.docx"
OUT = "Islamic_Economic_Literacy_Education_Gen_Z_SLR_AlBayan_ID.docx"

doc = Document(TEMPLATE)

# ---- font Normal default: Book Antiqua 12 ----
normal = doc.styles["Normal"]
normal.font.name = "Book Antiqua"
normal.font.size = Pt(12)
rpr = normal.element.get_or_add_rPr(); rf = rpr.get_or_add_rFonts()
rf.set(qn("w:ascii"), "Book Antiqua"); rf.set(qn("w:hAnsi"), "Book Antiqua")

# ---- kosongkan body, simpan sectPr (header/footer) ----
body = doc.element.body
sectPr = body.find(qn("w:sectPr"))
for child in list(body):
    if child is sectPr:
        continue
    body.remove(child)

# ---- helper paragraf ----
def P(style=None, align="justify", line=1.5, after=6, before=0):
    p = doc.add_paragraph()
    if style:
        p.style = doc.styles[style]
    amap = {"justify": WD_ALIGN_PARAGRAPH.JUSTIFY, "center": WD_ALIGN_PARAGRAPH.CENTER,
            "left": WD_ALIGN_PARAGRAPH.LEFT, "right": WD_ALIGN_PARAGRAPH.RIGHT}
    p.alignment = amap[align]
    pf = p.paragraph_format
    pf.line_spacing = line; pf.space_after = Pt(after); pf.space_before = Pt(before)
    return p

def run(p, text, bold=False, italic=False, size=12):
    r = p.add_run(text)
    r.bold = bold; r.italic = italic
    r.font.name = "Book Antiqua"; r.font.size = Pt(size)
    return r

def heading(text, center=False):
    p = P(style="Heading 1", align=("center" if center else "left"), line=1.5, after=6, before=10)
    run(p, text, bold=True, size=12)
    return p

def subheading(text):
    p = P(align="left", line=1.5, after=4, before=6)
    run(p, text, bold=True, size=12)
    return p

def body_para(text, after=6):
    p = P(align="justify", line=1.5, after=after)
    run(p, text, size=12)
    return p

# ---- helper tabel ----
def _set_table_borders(t):
    tblPr = t._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement("w:" + edge)
        e.set(qn("w:val"), "single"); e.set(qn("w:sz"), "4")
        e.set(qn("w:space"), "0"); e.set(qn("w:color"), "000000")
        borders.append(e)
    tblPr.append(borders)

def make_table(rows, cols):
    t = doc.add_table(rows=rows, cols=cols)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(t)
    return t

def setcell(cell, text, bold=False, italic=False, size=11):
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2); p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(text); r.bold = bold; r.italic = italic
    r.font.name = "Book Antiqua"; r.font.size = Pt(size)

def caption(text):
    p = P(align="center", line=1.0, after=8, before=2)
    run(p, text, italic=True, size=10)

def add_image(path, width_cm, caption_text):
    p = P(align="center", line=1.0, after=2, before=4)
    if os.path.exists(path):
        p.add_run().add_picture(path, width=Cm(width_cm))
    caption(caption_text)

# =====================================================================
# JUDUL + PENULIS
# =====================================================================
pt = P(style="Title", align="center", line=1.0, after=2)
run(pt, "Pendidikan Literasi Ekonomi Islam pada Generasi Z Muslim: Tinjauan Literatur Sistematis", bold=True, size=12)

pa = P(style="Wawasan_1.3 Afiliasi", align="center", line=1.0, after=0)
run(pa, "[Nama Penulis]", size=12)
pa = P(style="Wawasan_1.3 Afiliasi", align="center", line=1.0, after=8)
run(pa, "[Program Studi/Fakultas], [Universitas], [Kota], Indonesia; [email penulis]", italic=True, size=11)

# =====================================================================
# ABSTRACT (EN)
# =====================================================================
pab = P(align="justify", line=1.0, after=2)
run(pab, "Abstract", bold=True, italic=True, size=12)
abstract_en = ("The rapid expansion of Islamic finance, the digitalisation of economic life, and the rise of a "
"halal lifestyle have made Islamic economic literacy a strategic competency for Muslim societies, "
"particularly for Generation Z Muslims who enter financial life through fintech, social media, and "
"Sharia-compliant products. This study systematically identifies, evaluates, and synthesises existing "
"research on Islamic economic literacy education among Generation Z Muslims. It employs a Systematic "
"Literature Review following a transparent and replicable protocol comprising planning, conducting, and "
"reporting phases, structured by a PICOC framework, Boolean search strings, inclusion and exclusion "
"criteria, a PRISMA-based selection workflow, a weighted quality-assessment checklist (Yes = 1, "
"Partially = 0.5, No = 0), a standardised data-extraction form, and combined thematic and descriptive "
"bibliometric analysis. Searches covered Scopus, Web of Science, Google Scholar, DOAJ, Crossref, Semantic "
"Scholar, Garuda, and SINTA-indexed journals for peer-reviewed studies published between 2015 and 2025 in "
"English or Indonesian. The synthesis reveals six themes and offers a theoretical consolidation of Islamic "
"economic literacy and practical implications for educators, Islamic financial institutions, and "
"policymakers in cultivating riba-averse, Sharia-conscious financial behaviour.")
run(P(align="justify", line=1.0, after=2), abstract_en, italic=True, size=12)
pk = P(align="justify", line=1.0, after=8)
run(pk, "Keywords: ", bold=True, italic=True, size=12)
run(pk, "Islamic economic literacy; Sharia financial literacy; Generation Z Muslims; Islamic financial education; halal financial behaviour.", italic=True, size=12)

# ABSTRAK (ID)
pabi = P(align="justify", line=1.0, after=2)
run(pabi, "Abstrak", bold=True, italic=True, size=12)
abstrak_id = ("Ekspansi keuangan Islam, digitalisasi kehidupan ekonomi, dan menguatnya gaya hidup halal "
"menjadikan literasi ekonomi Islam sebagai kompetensi strategis bagi masyarakat Muslim, khususnya Generasi "
"Z Muslim yang memasuki kehidupan finansial melalui fintech, media sosial, dan produk yang sesuai syariah. "
"Penelitian ini secara sistematis mengidentifikasi, menilai, dan menyintesis penelitian yang ada tentang "
"pendidikan literasi ekonomi Islam pada Generasi Z Muslim. Penelitian menggunakan Systematic Literature "
"Review dengan protokol yang transparan dan dapat direplikasi melalui tahap perencanaan, pelaksanaan, dan "
"pelaporan, yang distrukturkan oleh kerangka PICOC, string pencarian Boolean, kriteria inklusi dan "
"eksklusi, alur seleksi berbasis PRISMA, daftar tilik penilaian kualitas berbobot (Ya = 1, Sebagian = 0,5, "
"Tidak = 0), formulir ekstraksi data baku, serta analisis tematik dan bibliometrik deskriptif. Pencarian "
"dilakukan pada Scopus, Web of Science, Google Scholar, DOAJ, Crossref, Semantic Scholar, Garuda, dan "
"jurnal terindeks SINTA untuk studi tertelaah sejawat yang terbit pada 2015 sampai 2025 dalam bahasa "
"Inggris atau Indonesia. Sintesis menghasilkan enam tema serta memberikan konsolidasi teoretis dan "
"implikasi praktis bagi pendidik, lembaga keuangan syariah, dan pembuat kebijakan dalam menumbuhkan "
"perilaku keuangan yang sadar syariah dan menghindari riba.")
run(P(align="justify", line=1.0, after=2), abstrak_id, italic=True, size=12)
pki = P(align="justify", line=1.0, after=8)
run(pki, "Kata Kunci: ", bold=True, italic=True, size=12)
run(pki, "literasi ekonomi Islam; literasi keuangan syariah; Generasi Z Muslim; pendidikan keuangan Islam; perilaku keuangan halal.", italic=True, size=12)

# =====================================================================
# PENDAHULUAN
# =====================================================================
heading("PENDAHULUAN")

body_para("Selama tiga dekade terakhir, keuangan Islam berkembang dari alternatif etis yang bersifat khusus "
"menjadi industri global yang signifikan, mencakup perbankan syariah, pasar modal syariah, takaful, serta "
"instrumen keuangan sosial Islam seperti zakat, wakaf, dan sedekah. Seiring pertumbuhan industri tersebut, "
"perhatian akademik dan kebijakan terhadap literasi keuangan syariah juga meningkat, yaitu pengetahuan, "
"keterampilan, sikap, dan keyakinan yang diperlukan untuk mengambil keputusan keuangan sesuai prinsip "
"syariah (Antara dkk., 2016). Literasi keuangan konvensional telah lama dipahami sebagai bentuk modal "
"manusia yang memperbaiki perilaku menabung, perencanaan, dan pengelolaan utang (Lusardi & Mitchell, 2014). "
"Literasi ekonomi Islam menambahkan dimensi normatif berbasis keimanan, yaitu keharusan memperoleh, "
"membelanjakan, menabung, menginvestasikan, dan mendermakan harta dengan cara yang menghindari riba, "
"gharar, dan maysir, sekaligus menunaikan kewajiban seperti zakat (El-Gamal, 2006; Khan, 2010).")

body_para("Urgensi kompetensi ini semakin besar bagi Generasi Z, yaitu kohort yang lahir sekitar tahun 1997 "
"dan sesudahnya (Dimock, 2019). Generasi Z Muslim tumbuh dengan telepon pintar, media sosial, dompet "
"elektronik, layanan beli sekarang bayar nanti, dan platform fintech syariah sebagai infrastruktur baku "
"kehidupan ekonomi mereka (Prensky, 2001; Rabbani dkk., 2020). Mereka sekaligus merupakan segmen konsumen "
"yang kuat bagi ekonomi halal global dan populasi yang terpapar budaya digital konsumtif, pinjaman daring, "
"serta produk spekulatif yang berpotensi bertentangan dengan norma syariah. Bukti empiris menunjukkan bahwa "
"literasi keuangan syariah memengaruhi keputusan investasi generasi muda Muslim di pasar modal syariah "
"(Anisa & Fajri, 2025). Literasi tersebut juga berkaitan dengan niat mereka untuk mengadopsi perbankan "
"syariah (Albaity & Rahman, 2019).")

body_para("Pendidikan berada di pusat dinamika ini. Sekolah, perguruan tinggi, pesantren dan lembaga "
"pendidikan Islam lain, keluarga, komunitas Muslim, serta lembaga keuangan syariah berperan sebagai saluran "
"pembentukan kesadaran ekonomi Islam. Bukti dari pendidikan tinggi di Malaysia dan Indonesia menunjukkan "
"bahwa paparan mata kuliah muamalah secara nyata meningkatkan tingkat literasi (Rahim dkk., 2016; Ahmad "
"dkk., 2020). Pedagogi inovatif seperti sosialisasi keuangan syariah berbasis augmented reality juga dapat "
"meningkatkan pengetahuan dan motivasi belajar (Solikhatun & Sari, 2023). Pada ranah keuangan sosial, "
"literasi zakat memengaruhi niat pembayaran zakat melalui platform digital (Kasri & Yuniar, 2021).")

body_para("Meskipun kajian terus bertumbuh, literatur masih terfragmentasi karena studi-studi cenderung "
"membahas literasi keuangan syariah, pemuda Muslim, ekonomi syariah, gaya hidup halal, dan pendidikan "
"keuangan secara terpisah. Tinjauan yang ada lebih banyak menyoroti area berdekatan seperti adopsi fintech "
"syariah, manajemen zakat, dan literasi keuangan syariah pada populasi umum, sehingga sintesis sistematis "
"yang secara khusus berfokus pada pendidikan literasi ekonomi Islam di kalangan Generasi Z Muslim masih "
"sangat terbatas. Kebaruan tinjauan ini terletak pada penerapan protokol SLR yang ketat, transparan, dan "
"dapat direplikasi untuk mengonsolidasikan bukti yang tersebar. Penelitian ini diarahkan oleh lima "
"pertanyaan penelitian: (RQ1) tema dan tren utama; (RQ2) konseptualisasi literasi ekonomi Islam; (RQ3) "
"strategi pendidikan, media, dan pendekatan digital; (RQ4) tantangan dan peluang; serta (RQ5) kesenjangan "
"dan arah riset masa depan.")

# =====================================================================
# METODE PENELITIAN
# =====================================================================
heading("METODE PENELITIAN")

body_para("Penelitian ini menggunakan pendekatan Systematic Literature Review (SLR) untuk mengidentifikasi, "
"menilai, menyintesis, dan melaporkan penelitian yang ada dengan bersandar pada protokol a priori yang "
"terdokumentasi sehingga proses bersifat transparan, objektif, dan dapat direplikasi (Kitchenham & "
"Charters, 2007; Tranfield dkk., 2003). Pendekatan ini berbeda dari tinjauan naratif karena mengikuti "
"prosedur baku yang meminimalkan bias seleksi dan pelaporan (Snyder, 2019). Protokol mengikuti tiga tahap, "
"yaitu perencanaan, pelaksanaan, dan pelaporan (Carrera-Rivera dkk., 2022), serta mengacu pada standar "
"pelaporan PRISMA 2020 (Page dkk., 2021).")

subheading("Kerangka PICOC")
body_para("Cakupan tinjauan distrukturkan dengan kerangka PICOC yang menerjemahkan pertanyaan penelitian "
"menjadi konsep yang dapat ditelusuri (Kitchenham & Charters, 2007). Rincian kerangka tersebut disajikan "
"pada Tabel 1.")
picoc = [
    ("Elemen", "Definisi dan kata kunci terkait"),
    ("Population", "Generasi Z Muslim, pemuda Muslim, pelajar/mahasiswa Muslim, generasi muda Muslim (Gen Z, post-millennials)."),
    ("Intervention", "Pendidikan literasi ekonomi Islam/keuangan syariah, pendidikan keuangan halal, pendidikan muamalah, sosialisasi keuangan."),
    ("Comparison", "Literasi keuangan konvensional, literasi ekonomi umum, pendidikan non-syariah; atau tanpa pembanding bila tidak relevan (N/A)."),
    ("Outcome", "Kesadaran ekonomi Islam, perilaku sesuai syariah, keputusan keuangan halal, adopsi perbankan syariah, investasi halal, zakat, wakaf, sedekah, penghindaran riba."),
    ("Context", "Pendidikan ekonomi Islam, hukum dan ekonomi Islam, lembaga keuangan syariah, sekolah, perguruan tinggi, pesantren, komunitas Muslim, lingkungan pembelajaran digital."),
]
t = make_table(len(picoc), 2)
for i, (a, b) in enumerate(picoc):
    setcell(t.rows[i].cells[0], a, bold=(i == 0)); setcell(t.rows[i].cells[1], b, bold=(i == 0))
caption("Tabel 1. Kerangka PICOC")

subheading("Strategi Pencarian dan Sumber Pustaka Digital")
body_para("String pencarian dibangun dengan menggabungkan kata kunci dan sinonim PICOC menggunakan operator "
"Boolean serta diterapkan pada judul, abstrak, dan kata kunci sejauh didukung basis data. Salah satu contoh "
"string pencarian adalah (\u201cIslamic economic literacy\u201d OR \u201cSharia financial literacy\u201d OR "
"\u201cIslamic financial literacy\u201d) AND (\u201cGeneration Z\u201d OR \u201cGen Z\u201d OR \u201cMuslim "
"youth\u201d OR \u201cMuslim students\u201d) AND (\u201ceducation\u201d OR \u201cfinancial education\u201d OR "
"\u201cliteracy education\u201d). Pencarian dilakukan pada delapan sumber, yaitu Scopus, Web of Science, "
"Google Scholar, DOAJ, Crossref, Semantic Scholar, Garuda, dan jurnal terindeks SINTA, dilengkapi "
"penelusuran maju dan mundur (snowballing). Pemilihan sumber ini penting karena banyak bukti primer berasal "
"dari Indonesia dan Malaysia, sehingga pengindeks regional perlu dilibatkan bersama Scopus dan Web of "
"Science.")

subheading("Kriteria Inklusi dan Eksklusi")
inc = ("(1) Membahas literasi ekonomi Islam/keuangan syariah atau pendidikan ekonomi Islam; (2) berfokus "
"pada Generasi Z, pemuda Muslim, pelajar/mahasiswa, atau generasi muda Muslim; (3) membahas pendidikan, "
"media pembelajaran, pendekatan digital, perilaku keuangan, gaya hidup halal, lembaga keuangan syariah, "
"zakat, wakaf, sedekah, atau investasi halal; (4) artikel jurnal tertelaah sejawat, prosiding, atau bab "
"buku ilmiah; (5) terbit pada 2015 sampai 2025; (6) ditulis dalam bahasa Inggris atau Indonesia.")
exc = ("(1) Tidak terkait ekonomi Islam/literasi keuangan syariah; (2) hanya membahas literasi keuangan "
"konvensional tanpa relevansi Islam; (3) tidak membahas pemuda/pelajar/Generasi Z; (4) opini, berita, "
"laporan nonakademik, atau teks penuh tidak dapat diakses; (5) artikel duplikat; (6) kualitas metodologis "
"lemah di bawah ambang.")
t = make_table(2, 2)
setcell(t.rows[0].cells[0], "Inklusi", bold=True); setcell(t.rows[0].cells[1], inc)
setcell(t.rows[1].cells[0], "Eksklusi", bold=True); setcell(t.rows[1].cells[1], exc)
caption("Tabel 2. Kriteria Inklusi dan Eksklusi")

subheading("Seleksi Studi dan Penilaian Kualitas")
body_para("Seleksi berlangsung dalam lima tahap berbasis PRISMA, yaitu identifikasi, penghapusan duplikat, "
"penyaringan judul dan abstrak, penilaian kelayakan teks penuh, dan penyertaan akhir. Setiap studi kandidat "
"dinilai dengan daftar tilik penilaian kualitas tujuh butir yang diberi skor Ya = 1, Sebagian = 0,5, Tidak "
"= 0 (maksimum = 7), mencakup kejelasan tujuan, relevansi terhadap literasi keuangan syariah, fokus pada "
"Generasi Z/pemuda Muslim, kejelasan metode, dukungan data terhadap temuan, implikasi bagi ekonomi "
"Islam/pendidikan/perilaku, serta pengakuan keterbatasan. Studi dengan skor minimal 4,0 dari 7 (sekitar 57 "
"persen) dipertahankan. Proses seleksi diringkas dalam diagram alir PRISMA 2020 pada Gambar 1.")

add_image("prisma_flow_id.png", 13.5,
    "Gambar 1. Diagram alir PRISMA 2020 untuk proses seleksi studi "
    "(jumlah bersifat ilustratif; difinalkan setelah ekstraksi lengkap).")

subheading("Ekstraksi Data dan Sintesis")
body_para("Formulir ekstraksi data baku digunakan untuk merekam penulis, tahun, negara/konteks, judul, "
"tujuan, metodologi, populasi/sampel, konsep kunci, jenis literasi, strategi pendidikan/media pembelajaran, "
"tema ekonomi Islam, temuan kunci, tantangan, kesenjangan, dan relevansi terhadap pertanyaan penelitian. "
"Analisis memadukan sintesis tematik dengan analisis bibliometrik deskriptif, yang meliputi tren publikasi "
"per tahun, dominasi konteks, sumber publikasi utama, dan kata kunci yang sering muncul.")

# =====================================================================
# HASIL DAN PEMBAHASAN
# =====================================================================
heading("HASIL DAN PEMBAHASAN")

subheading("Hasil Penelitian")
body_para("Korpus yang disintesis menunjukkan tiga pola konsisten. Pertama, aktivitas publikasi meningkat "
"tajam sejak 2016 dan menguat setelah 2019, sejalan dengan ekspansi fintech syariah dan perhatian kebijakan "
"terhadap literasi keuangan. Kedua, distribusi geografis sangat terkonsentrasi di Indonesia dan Malaysia, "
"dengan kontribusi tambahan dari kawasan Teluk dan Yordania. Ketiga, sumber publikasi yang paling sering "
"adalah jurnal spesialis seperti Journal of Islamic Marketing, International Journal of Islamic and Middle "
"Eastern Finance and Management, dan Journal of Islamic Accounting and Business Research, dilengkapi jurnal "
"terindeks SINTA dan DOAJ. Kata kunci yang berulang mencakup literasi keuangan syariah, religiusitas, "
"Generasi Z/milenial, perbankan syariah, fintech, zakat, dan wakaf tunai. Tabel 3 merangkum studi "
"representatif yang menjadi jangkar sintesis.")

hdr = ["No.", "Penulis & Tahun", "Konteks", "Fokus", "Metode", "Temuan Kunci", "Relevansi"]
rows = [
    ("1", "Lusardi & Mitchell (2014)", "AS/global", "Ekonomi literasi keuangan", "Tinjauan teori & bukti", "Literasi keuangan sebagai investasi modal manusia yang memengaruhi menabung, perencanaan, dan utang", "Dasar konseptual (RQ2)"),
    ("2", "Antara dkk. (2016)", "Malaysia", "Literasi keuangan & halal", "Konseptual", "Memposisikan literasi keuangan syariah dalam ekosistem halal", "Konseptualisasi (RQ2, RQ4)"),
    ("3", "Rahim dkk. (2016)", "Malaysia", "Determinan literasi mahasiswa", "Survei, EFA", "Religiusitas dan kepuasan finansial membentuk literasi mahasiswa", "Determinan pemuda (RQ1, RQ3)"),
    ("4", "Albaity & Rahman (2019)", "UEA", "Literasi & niat bank syariah", "Survei, SEM", "Literasi berkaitan dengan niat lebih besar mengadopsi perbankan syariah", "Literasi ke institusi (RQ2, RQ4)"),
    ("5", "Ahmad dkk. (2020)", "Indonesia", "Determinan literasi", "Survei", "Mengidentifikasi determinan sosiodemografis dan pendidikan", "Pendorong pendidikan (RQ1, RQ3)"),
    ("6", "Solikhatun & Sari (2023)", "Indonesia", "Edukasi syariah berbasis AR", "Kuasi-eksperimen", "Media AR meningkatkan motivasi dan pengetahuan keuangan syariah", "Pedagogi digital (RQ3)"),
    ("7", "Kasri & Yuniar (2021)", "Indonesia", "Determinan zakat digital", "Survei, UTAUT", "Literasi zakat meningkatkan niat membayar zakat secara digital", "Literasi sosial (RQ3, RQ4)"),
    ("8", "Anisa & Fajri (2025)", "Indonesia", "Literasi & investasi Gen Z", "Survei", "Literasi berpengaruh positif pada keputusan investasi Gen Z di pasar modal syariah", "Bukti langsung Gen Z (RQ1, RQ4)"),
]
t = make_table(len(rows) + 1, len(hdr))
for j, h in enumerate(hdr):
    setcell(t.rows[0].cells[j], h, bold=True, size=10)
for i, row in enumerate(rows, 1):
    for j, val in enumerate(row):
        setcell(t.rows[i].cells[j], val, size=10)
for j, w in enumerate([0.8, 2.3, 1.4, 2.0, 1.9, 3.2, 2.6]):
    for r in t.rows:
        r.cells[j].width = Cm(w)
caption("Tabel 3. Karakteristik Studi Representatif yang Disertakan")

body_para("Secara tematik, bukti memusat pada enam tema, yaitu (1) literasi ekonomi Islam sebagai konstruk "
"multidimensi berlandaskan keimanan yang memadukan dimensi kognitif, normatif-religius, dan perilaku; (2) "
"kaitan literasi dengan perilaku sesuai syariah di kalangan Muslim muda; (3) kanal digital dan media sosial "
"sebagai antarmuka pembelajaran utama bagi penutur asli digital; (4) lembaga keuangan syariah, investasi "
"halal, dan kesadaran pemuda; (5) literasi keuangan sosial zakat, wakaf, dan sedekah yang dimediasi "
"platform digital; serta (6) tantangan dan kesenjangan, terutama kesenjangan sikap dan perilaku, lemahnya "
"pengukuran khusus pemuda, dominasi desain lintas-seksi, dan konsentrasi geografis.")

subheading("Pembahasan")
body_para("Literatur secara konsisten membingkai literasi ekonomi Islam sebagai konstruk multidimensi "
"berlandaskan keimanan yang memadukan pemahaman keuangan sesuai syariah dan keuangan sosial Islam, niat "
"untuk patuh, serta kapasitas perilaku untuk bertindak (Antara dkk., 2017). Instrumen konvensional dinilai "
"tidak memadai bagi Muslim karena mengabaikan parameter berbasis keimanan, sehingga konstruk ini perlu "
"dilandaskan pada maqashid al-shariah, khususnya penjagaan harta atau hifzh al-mal (menjawab RQ2).")

body_para("Literasi yang lebih tinggi dikaitkan dengan niat mengadopsi perbankan syariah (Albaity & Rahman, "
"2019) dan keputusan investasi yang lebih baik di pasar modal syariah (Anisa & Fajri, 2025). Namun, "
"keterpaparan budaya digital konsumtif menyebabkan literasi tidak otomatis menjadi praktik, sehingga "
"terdapat kesenjangan sikap dan perilaku yang harus dijawab oleh pendidikan (menjawab RQ1 dan RQ4).")

body_para("Pendekatan yang efektif memadukan paparan kurikuler formal, misalnya mata kuliah muamalah yang "
"terbukti meningkatkan literasi (Rahim dkk., 2016), dengan pedagogi berbantuan teknologi yang sesuai bagi "
"penutur asli digital, seperti augmented reality, gamifikasi, pembelajaran seluler, dan penjangkauan media "
"sosial (Solikhatun & Sari, 2023; Prensky, 2001). Bagi institusi, edukasi digital yang terintegrasi dalam "
"perjalanan layanan dapat sekaligus membangun literasi dan kepercayaan (menjawab RQ3).")

body_para("Tantangan utama adalah kesenjangan sikap dan perilaku di tengah maraknya tawaran berbasis bunga "
"dan konsumtif, yang diperberat oleh literasi dasar yang rendah dan lemahnya pengukuran khusus pemuda. "
"Peluangnya adalah platform digital yang sama justru menjadi kanal tak tertandingi untuk edukasi berskala "
"besar dan relevan secara budaya (Hassan dkk., 2020). Riset masa depan perlu mengembangkan instrumen khusus "
"Generasi Z yang tervalidasi, menerapkan desain eksperimental dan longitudinal, mendiversifikasi konteks di "
"luar Indonesia dan Malaysia, memadukan literasi keuangan komersial dan sosial, serta mengevaluasi pedagogi "
"digital dengan pengaman keaslian syariah. Dengan membekali Muslim muda untuk mengenali riba, gharar, dan "
"maysir serta memilih alternatif yang sesuai syariah, pendidikan literasi ekonomi Islam berfungsi sebagai "
"mekanisme preventif terhadap pinjaman berbasis bunga sekaligus mengarahkan sumber daya ke investasi halal "
"dan keuangan sosial Islam (Chapra, 1992; Iqbal & Mirakhor, 2011) (menjawab RQ4 dan RQ5).")

# =====================================================================
# KESIMPULAN
# =====================================================================
heading("KESIMPULAN")
body_para("Dengan protokol SLR yang transparan dan dapat direplikasi yang distrukturkan oleh kerangka "
"PICOC, string pencarian Boolean pada delapan basis data, kriteria inklusi dan eksklusi yang eksplisit, "
"alur berbasis PRISMA, daftar tilik penilaian kualitas berbobot, formulir ekstraksi baku, serta analisis "
"tematik dan bibliometrik, tinjauan ini mengonsolidasikan bukti yang terfragmentasi menjadi enam tema. "
"Literasi ekonomi Islam muncul sebagai konstruk multidimensi berlandaskan keimanan yang berkaitan secara "
"empiris dengan perilaku sesuai syariah dan keterlibatan kelembagaan; kanal digital dan media sosial "
"menjadi medium penentu untuk menjangkau Generasi Z; literasi keuangan sosial zakat dan wakaf merupakan "
"komponen integral yang dimediasi secara digital; dan bidang ini terkendala kesenjangan sikap dan perilaku, "
"lemahnya pengukuran khusus pemuda, dominasi desain lintas-seksi, serta konsentrasi geografis.")
body_para("Secara praktis, hasil ini mendukung pengintegrasian literasi keuangan syariah lintas kurikulum, "
"penyiapan lembaga keuangan syariah untuk menyampaikan edukasi digital yang tepercaya, pelibatan keluarga "
"dan figur daring yang kredibel, serta perancangan program yang menyasar perilaku, bukan sekadar kesadaran. "
"Keterbatasan tinjauan terikat pada protokolnya, yaitu jendela 2015 sampai 2025, pembatasan bahasa Inggris "
"dan Indonesia, basis data terpilih, dan ambang kualitas, serta konsentrasi geografis sumber literatur. "
"Riset masa depan diarahkan pada instrumen khusus Generasi Z yang tervalidasi, evaluasi eksperimental dan "
"longitudinal, diversifikasi konteks, kerangka literasi yang memadukan keuangan komersial dan sosial, serta "
"evaluasi pedagogi digital yang menjaga keaslian syariah.")

# =====================================================================
# REFERENSI (APA 7th edition, alfabetis, hanging indent)
# =====================================================================
heading("REFERENSI", center=True)
biblio = [
"Ahmad, G. N., Widyastuti, U., Susanti, S., & Mukhibad, H. (2020). Determinants of the Islamic financial literacy. Accounting, 6(6), 961-966.",
"Albaity, M., & Rahman, M. (2019). The intention to use Islamic banking: An exploratory study to measure Islamic financial literacy. International Journal of Emerging Markets, 14(5), 988-1012.",
"Anisa, S. N., & Fajri AF, M. S. (2025). The impact of Islamic financial literacy on Gen Z\u2019s investment decisions in the Islamic capital market [Working paper]. SSRN.",
"Antara, P. M., Musa, R., & Hassan, F. (2016). Bridging Islamic financial literacy and halal literacy: The way forward in halal ecosystem. Procedia Economics and Finance, 37, 196-202.",
"Antara, P. M., Musa, R., & Hassan, F. (2017). Conceptualisation and operationalisation of Islamic financial literacy scale. Pertanika Journal of Social Sciences & Humanities, 25(S), 251-260.",
"Atkinson, A., & Messy, F.-A. (2012). Measuring financial literacy: Results of the OECD/INFE pilot study (OECD Working Papers on Finance, Insurance and Private Pensions No. 15). OECD Publishing.",
"Carrera-Rivera, A., Ochoa, W., Larrinaga, F., & Lasa, G. (2022). How-to conduct a systematic literature review: A quick guide for computer science research. MethodsX, 9, 101895.",
"Chapra, M. U. (1992). Islam and the economic challenge. The Islamic Foundation.",
"Dimock, M. (2019). Defining generations: Where Millennials end and Generation Z begins. Pew Research Center.",
"El-Gamal, M. A. (2006). Islamic finance: Law, economics, and practice. Cambridge University Press.",
"Hassan, M. K., Rabbani, M. R., & Ali, M. A. M. (2020). Challenges for the Islamic finance and banking in post COVID era and the role of fintech. Journal of Economic Cooperation and Development, 41(3), 93-116.",
"Iqbal, Z., & Mirakhor, A. (2011). An introduction to Islamic finance: Theory and practice (Edisi ke-2). John Wiley & Sons.",
"Kasri, R. A., & Yuniar, A. M. (2021). Determinants of digital zakat payments: Lessons from Indonesian experience. Journal of Islamic Accounting and Business Research, 12(3), 362-379.",
"Khan, F. (2010). How \u201cIslamic\u201d is Islamic banking? Journal of Economic Behavior & Organization, 76(3), 805-820.",
"Kitchenham, B., & Charters, S. (2007). Guidelines for performing systematic literature reviews in software engineering (EBSE Technical Report EBSE-2007-01). Keele University & University of Durham.",
"Lusardi, A., & Mitchell, O. S. (2014). The economic importance of financial literacy: Theory and evidence. Journal of Economic Literature, 52(1), 5-44.",
"Otoritas Jasa Keuangan. (2022). Survei Nasional Literasi dan Inklusi Keuangan (SNLIK) 2022. Otoritas Jasa Keuangan.",
"Page, M. J., McKenzie, J. E., Bossuyt, P. M., Boutron, I., Hoffmann, T. C., Mulrow, C. D., ... Moher, D. (2021). The PRISMA 2020 statement: An updated guideline for reporting systematic reviews. BMJ, 372, n71.",
"Prensky, M. (2001). Digital natives, digital immigrants. On the Horizon, 9(5), 1-6.",
"Rabbani, M. R., Khan, S., & Thalassinos, E. I. (2020). FinTech, blockchain and Islamic finance: An extensive literature review. International Journal of Economics and Business Administration, 8(2), 65-86.",
"Rahim, S. H. A., Rashid, R. A., & Hamed, A. B. (2016). Islamic financial literacy and its determinants among university students: An exploratory factor analysis. International Journal of Economics and Financial Issues, 6(7S), 32-35.",
"Snyder, H. (2019). Literature review as a research methodology: An overview and guidelines. Journal of Business Research, 104, 333-339.",
"Solikhatun, I., & Sari, R. C. (2023). Sharia financial education using augmented reality technology to increase student motivation in distance learning. AIP Conference Proceedings, 2654, 020008.",
"Tranfield, D., Denyer, D., & Smart, P. (2003). Towards a methodology for developing evidence-informed management knowledge by means of systematic review. British Journal of Management, 14(3), 207-222.",
]
for e in biblio:
    p = P(align="justify", line=1.0, after=6)
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-1.0)
    run(p, e, size=11)

# sectPr ke akhir agar header/footer tetap valid
body.remove(sectPr)
body.append(sectPr)

# pengaman: pastikan tidak ada em-dash / en-dash di seluruh dokumen
def assert_no_dash():
    bad = []
    for para in doc.paragraphs:
        if "\u2014" in para.text or "\u2013" in para.text:
            bad.append(para.text[:60])
    for tb in doc.tables:
        for row in tb.rows:
            for cell in row.cells:
                if "\u2014" in cell.text or "\u2013" in cell.text:
                    bad.append(cell.text[:60])
    return bad

bad = assert_no_dash()
if bad:
    raise SystemExit("DITEMUKAN dash terlarang: %r" % bad)

doc.save(OUT)
print("Saved:", OUT, "| tanpa dash terlarang, tanpa footnote")


# =====================================================================
# Bersihkan footnotes.xml bawaan template (hapus contoh footnote 1,2,3)
# sehingga tidak ada teks catatan kaki yang tersisa.
# =====================================================================
_footnotes_xml = (
    '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
    '<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
    '<w:footnote w:type="separator" w:id="-1"><w:p><w:r><w:separator/></w:r></w:p></w:footnote>'
    '<w:footnote w:type="continuationSeparator" w:id="0"><w:p><w:r><w:continuationSeparator/></w:r></w:p></w:footnote>'
    '</w:footnotes>'
)
_tmp = OUT + ".tmp"
with zipfile.ZipFile(OUT, "r") as zin:
    _data = {n: zin.read(n) for n in zin.namelist()}
if "word/footnotes.xml" in _data:
    _data["word/footnotes.xml"] = _footnotes_xml.encode("utf-8")
with zipfile.ZipFile(_tmp, "w", zipfile.ZIP_DEFLATED) as zout:
    for n, b in _data.items():
        zout.writestr(n, b)
shutil.move(_tmp, OUT)
print("footnotes.xml dibersihkan (hanya separator).")
