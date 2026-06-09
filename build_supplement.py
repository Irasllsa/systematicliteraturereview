# -*- coding: utf-8 -*-
"""
Membuat berkas Materi Suplemen (data dukung) untuk artikel SLR
yang diajukan ke Al-Bayan: Journal of Islamic Law and Economics.
Bahasa Indonesia, tanpa catatan kaki, tanpa em-dash/en-dash.
"""
import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "Al-Bayan_Supplementary_Material_SLR.docx"

doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "Book Antiqua"
normal.font.size = Pt(11)
rpr = normal.element.get_or_add_rPr(); rf = rpr.get_or_add_rFonts()
rf.set(qn("w:ascii"), "Book Antiqua"); rf.set(qn("w:hAnsi"), "Book Antiqua")

for s in doc.sections:
    s.page_height = Cm(29.7); s.page_width = Cm(21.0)
    s.left_margin = Cm(2.5); s.right_margin = Cm(2.5)
    s.top_margin = Cm(2.5); s.bottom_margin = Cm(2.5)

def P(align="justify", line=1.0, after=6, before=0):
    p = doc.add_paragraph()
    amap = {"justify": WD_ALIGN_PARAGRAPH.JUSTIFY, "center": WD_ALIGN_PARAGRAPH.CENTER,
            "left": WD_ALIGN_PARAGRAPH.LEFT}
    p.alignment = amap[align]
    pf = p.paragraph_format
    pf.line_spacing = line; pf.space_after = Pt(after); pf.space_before = Pt(before)
    return p

def run(p, text, bold=False, italic=False, size=11):
    r = p.add_run(text); r.bold = bold; r.italic = italic
    r.font.name = "Book Antiqua"; r.font.size = Pt(size)
    return r

def h1(text):
    p = P(align="left", after=6, before=12); run(p, text, bold=True, size=13); return p

def h2(text):
    p = P(align="left", after=4, before=8); run(p, text, bold=True, size=12); return p

def body(text, after=6):
    p = P(after=after); run(p, text, size=11); return p

def _borders(t):
    tblPr = t._tbl.tblPr
    b = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement("w:" + edge)
        e.set(qn("w:val"), "single"); e.set(qn("w:sz"), "4")
        e.set(qn("w:space"), "0"); e.set(qn("w:color"), "000000")
        b.append(e)
    tblPr.append(b)

def table(rows, cols):
    t = doc.add_table(rows=rows, cols=cols)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    _borders(t)
    return t

def cell(c, text, bold=False, size=9, align="left"):
    p = c.paragraphs[0]
    p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER}[align]
    p.paragraph_format.space_after = Pt(1); p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(text); r.bold = bold; r.font.name = "Book Antiqua"; r.font.size = Pt(size)

def caption(text):
    p = P(align="center", after=8, before=2); run(p, text, italic=True, size=9)

# HEADER
p = P(align="center", after=2); run(p, "MATERI SUPLEMEN (DATA DUKUNG)", bold=True, size=14)
p = P(align="center", after=2); run(p, "Pendidikan Literasi Ekonomi Islam pada Generasi Z Muslim: Tinjauan Literatur Sistematis", bold=True, size=12)
p = P(align="center", after=2); run(p, "Diajukan ke Al-Bayan: Journal of Islamic Law and Economics", italic=True, size=11)
p = P(align="center", after=10); run(p, "Penulis: [Nama Penulis]  |  Korespondensi: [email]", size=10)

body("Berkas ini memuat materi suplemen yang mendukung keterbukaan dan keterulangan (reproducibility) "
"penelitian, yaitu (A) strategi pencarian per basis data, (B) daftar tilik PRISMA 2020, (C) matriks "
"penilaian kualitas studi, (D) formulir ekstraksi data lengkap, dan (E) diagram alir PRISMA 2020. Korpus "
"yang ditampilkan merupakan korpus representatif yang menjadi jangkar sintesis; matriks dan ekstraksi "
"final difinalkan setelah pencarian dijalankan ulang pada tanggal pengajuan untuk memastikan jumlah "
"rekaman terbaru.")

# LAMPIRAN A
h1("Lampiran A. Strategi Pencarian dan Sumber Basis Data")
body("Pencarian diterapkan pada judul, abstrak, dan kata kunci (jika didukung), pada rentang terbit 2015 "
"sampai 2025, dalam bahasa Inggris atau Indonesia, dan dilengkapi penelusuran maju dan mundur (snowballing). "
"Tanggal pelaksanaan pencarian: [isi tanggal, mis. 9 Juni 2026].")
body("String pencarian inti (disesuaikan dengan sintaks tiap basis data):", after=2)
body("String 1: (\u201cIslamic economic literacy\u201d OR \u201cSharia economic literacy\u201d OR \u201cIslamic "
"financial literacy\u201d OR \u201cSharia financial literacy\u201d) AND (\u201cGeneration Z\u201d OR \u201cGen "
"Z\u201d OR \u201cMuslim youth\u201d OR \u201cyoung Muslims\u201d OR \u201cMuslim students\u201d) AND "
"(\u201ceducation\u201d OR \u201clearning\u201d OR \u201cfinancial education\u201d OR \u201cliteracy "
"education\u201d).", after=2)
body("String 2: (\u201cIslamic financial literacy education\u201d OR \u201cSharia financial education\u201d) AND "
"(\u201cMuslim students\u201d OR \u201cGeneration Z Muslims\u201d) AND (\u201chalal financial behavior\u201d OR "
"\u201cSharia-compliant behavior\u201d OR \u201cIslamic economic awareness\u201d).", after=2)
body("String 3: (\u201chalal financial literacy\u201d OR \u201cIslamic economic awareness\u201d) AND "
"(\u201cdigital media\u201d OR \u201csocial media\u201d OR \u201conline learning\u201d OR \u201cfinancial "
"technology\u201d) AND (\u201cyoung Muslims\u201d OR \u201cGeneration Z\u201d).", after=6)

rowsA = [
    ("No.", "Basis data", "Ruang lingkup pencarian", "Jumlah rekaman"),
    ("1", "Scopus", "TITLE-ABS-KEY( ... )", "[isi]"),
    ("2", "Web of Science", "Topic search TS=( ... )", "[isi]"),
    ("3", "Google Scholar", "Seluruh rekaman; disaring sampai titik jenuh", "[isi]"),
    ("4", "DOAJ", "Judul, abstrak, kata kunci", "[isi]"),
    ("5", "Crossref", "Metadata judul dan DOI (verifikasi)", "[isi]"),
    ("6", "Semantic Scholar", "Penelusuran sitasi (snowballing)", "[isi]"),
    ("7", "Garuda", "Judul, abstrak, kata kunci (Indonesia)", "[isi]"),
    ("8", "Jurnal terindeks SINTA", "Judul, abstrak, kata kunci (Indonesia)", "[isi]"),
    ("", "Total (sebelum dedup)", "", "[isi]"),
]
t = table(len(rowsA), 4)
for i, r in enumerate(rowsA):
    for j, v in enumerate(r):
        cell(t.rows[i].cells[j], v, bold=(i == 0))
for j, w in enumerate([1.0, 4.0, 8.0, 2.5]):
    for r in t.rows:
        r.cells[j].width = Cm(w)
caption("Tabel A1. Sumber basis data, ruang lingkup pencarian, dan jumlah rekaman (diisi pada finalisasi).")

# LAMPIRAN B
h1("Lampiran B. Daftar Tilik PRISMA 2020")
body("Daftar tilik berikut mengacu pada PRISMA 2020 (Page dkk., 2021) dan menunjukkan lokasi pelaporan tiap "
"butir di dalam naskah. Butir yang tidak relevan ditandai sebagai tidak berlaku karena tinjauan ini bersifat "
"kualitatif-tematik dan tidak melakukan meta-analisis.")
prisma = [
    ("No.", "Butir (ringkas)", "Dilaporkan pada"),
    ("1", "Judul mengidentifikasi sebagai tinjauan sistematis", "Judul"),
    ("2", "Abstrak terstruktur", "Abstrak / Abstract"),
    ("3", "Rasionalisasi", "Pendahuluan"),
    ("4", "Tujuan dan pertanyaan penelitian", "Pendahuluan"),
    ("5", "Kriteria kelayakan", "Metode (Kriteria Inklusi dan Eksklusi; Tabel 2)"),
    ("6", "Sumber informasi", "Metode (Strategi Pencarian); Lampiran A"),
    ("7", "Strategi pencarian", "Metode; Lampiran A"),
    ("8", "Proses seleksi", "Metode (Seleksi Studi dan Penilaian Kualitas); Gambar 1"),
    ("9", "Proses pengumpulan data", "Metode (Ekstraksi Data dan Sintesis); Lampiran D"),
    ("10", "Item data", "Metode (Ekstraksi Data); Lampiran D"),
    ("11", "Penilaian risiko bias / kualitas studi", "Metode (Penilaian Kualitas); Lampiran C"),
    ("12", "Ukuran efek", "Tidak berlaku (sintesis kualitatif-tematik)"),
    ("13", "Metode sintesis", "Metode (Ekstraksi Data dan Sintesis)"),
    ("14", "Penilaian bias pelaporan", "Tidak berlaku"),
    ("15", "Penilaian kepastian bukti", "Dibobot melalui skor kualitas (Lampiran C)"),
    ("16", "Seleksi studi (hasil)", "Hasil dan Pembahasan; Gambar 1"),
    ("17", "Karakteristik studi", "Hasil dan Pembahasan (Tabel 3); Lampiran D"),
    ("18", "Risiko bias dalam studi", "Lampiran C"),
    ("19", "Hasil tiap studi", "Lampiran D"),
    ("20", "Hasil sintesis", "Hasil dan Pembahasan (Temuan tematik)"),
    ("21", "Bias pelaporan", "Tidak berlaku"),
    ("22", "Kepastian bukti", "Pembahasan"),
    ("23", "Diskusi (ringkasan, keterbatasan, implikasi)", "Pembahasan; Kesimpulan"),
    ("24", "Registrasi dan protokol", "Belum diregistrasi; protokol tersedia atas permintaan"),
    ("25", "Dukungan/pendanaan", "[isi: tidak ada pendanaan khusus / sebutkan]"),
    ("26", "Konflik kepentingan", "Tidak ada konflik kepentingan"),
    ("27", "Ketersediaan data dan materi", "Materi suplemen ini; berkas RIS referensi"),
]
t = table(len(prisma), 3)
for i, r in enumerate(prisma):
    for j, v in enumerate(r):
        cell(t.rows[i].cells[j], v, bold=(i == 0))
for j, w in enumerate([1.0, 8.0, 6.5]):
    for r in t.rows:
        r.cells[j].width = Cm(w)
caption("Tabel B1. Daftar tilik PRISMA 2020 dan lokasi pelaporan.")

# LAMPIRAN C
h1("Lampiran C. Matriks Penilaian Kualitas Studi")
body("Setiap studi dinilai pada tujuh butir dengan skor Ya = 1, Sebagian = 0,5, Tidak = 0 (maksimum 7). "
"Ambang inklusi adalah skor minimal 4,0. Keterangan butir: QA1 kejelasan tujuan; QA2 relevansi terhadap "
"literasi ekonomi/keuangan syariah; QA3 fokus pada Generasi Z, pemuda, atau mahasiswa Muslim; QA4 kejelasan "
"metode; QA5 dukungan data terhadap temuan; QA6 implikasi bagi ekonomi Islam, pendidikan, atau perilaku; "
"QA7 pengakuan keterbatasan dan arah riset. Skor bersifat ilustratif untuk korpus representatif.")
qa_hdr = ["No.", "Studi", "QA1", "QA2", "QA3", "QA4", "QA5", "QA6", "QA7", "Total", "Keputusan"]
qa_rows = [
    ("1", "Lusardi & Mitchell (2014)", "1", "0,5", "0", "1", "1", "0,5", "1", "5,0", "Rujukan konseptual"),
    ("2", "Antara dkk. (2016)", "1", "1", "0,5", "0,5", "0,5", "1", "0,5", "5,0", "Disertakan"),
    ("3", "Rahim dkk. (2016)", "1", "1", "1", "1", "1", "1", "0,5", "6,5", "Disertakan"),
    ("4", "Albaity & Rahman (2019)", "1", "1", "0,5", "1", "1", "1", "1", "6,5", "Disertakan"),
    ("5", "Ahmad dkk. (2020)", "1", "1", "0,5", "1", "1", "1", "0,5", "6,0", "Disertakan"),
    ("6", "Solikhatun & Sari (2023)", "1", "1", "1", "1", "1", "1", "0,5", "6,5", "Disertakan"),
    ("7", "Kasri & Yuniar (2021)", "1", "1", "0,5", "1", "1", "1", "1", "6,5", "Disertakan"),
    ("8", "Anisa & Fajri (2025)", "1", "1", "1", "0,5", "0,5", "1", "0,5", "5,5", "Disertakan"),
]
t = table(len(qa_rows) + 1, len(qa_hdr))
for j, v in enumerate(qa_hdr):
    cell(t.rows[0].cells[j], v, bold=True, size=9, align="center")
for i, r in enumerate(qa_rows, 1):
    for j, v in enumerate(r):
        cell(t.rows[i].cells[j], v, size=9, align=("left" if j == 1 else "center"))
widths = [0.8, 3.6, 0.95, 0.95, 0.95, 0.95, 0.95, 0.95, 0.95, 1.0, 2.6]
for j, w in enumerate(widths):
    for r in t.rows:
        r.cells[j].width = Cm(w)
caption("Tabel C1. Matriks skor penilaian kualitas (ambang inklusi: skor minimal 4,0 dari 7).")

# LAMPIRAN D
h1("Lampiran D. Formulir Ekstraksi Data Lengkap")
body("Setiap studi diekstraksi menggunakan formulir baku berikut. Untuk butir yang tidak dilaporkan secara "
"rinci di sini, pembaca dirujuk ke artikel asli.")
studies = [
    {"Penulis dan tahun": "Lusardi, A., & Mitchell, O. S. (2014)", "Negara/konteks": "Amerika Serikat / global",
     "Tujuan": "Mengkaji peran ekonomi literasi keuangan secara teoretis dan empiris", "Metodologi": "Tinjauan teori dan bukti",
     "Populasi/sampel": "Lintas populasi (kajian pustaka)", "Jenis literasi": "Literasi keuangan (konvensional)",
     "Strategi/media pendidikan": "Tidak spesifik", "Tema ekonomi Islam": "Tidak langsung; dasar konseptual",
     "Temuan kunci": "Literasi keuangan adalah investasi modal manusia yang memengaruhi menabung, perencanaan, dan utang",
     "Tantangan": "Pengukuran literasi dan endogenitas", "Relevansi RQ": "RQ2 (dasar konseptual)"},
    {"Penulis dan tahun": "Antara, P. M., Musa, R., & Hassan, F. (2016)", "Negara/konteks": "Malaysia",
     "Tujuan": "Menjembatani literasi keuangan syariah dan literasi halal", "Metodologi": "Konseptual",
     "Populasi/sampel": "Masyarakat Muslim (kerangka konsep)", "Jenis literasi": "Literasi keuangan syariah dan literasi halal",
     "Strategi/media pendidikan": "Penguatan ekosistem halal", "Tema ekonomi Islam": "Kepatuhan syariah, ekosistem halal",
     "Temuan kunci": "Literasi keuangan syariah menjadi fondasi partisipasi dalam ekosistem halal",
     "Tantangan": "Keterbatasan operasionalisasi konstruk", "Relevansi RQ": "RQ2, RQ4"},
    {"Penulis dan tahun": "Rahim, S. H. A., Rashid, R. A., & Hamed, A. B. (2016)", "Negara/konteks": "Malaysia",
     "Tujuan": "Mengidentifikasi determinan literasi keuangan syariah pada mahasiswa", "Metodologi": "Survei kuantitatif; analisis faktor eksploratori (EFA)",
     "Populasi/sampel": "Mahasiswa (lihat artikel asli untuk ukuran sampel)", "Jenis literasi": "Literasi keuangan syariah",
     "Strategi/media pendidikan": "Konteks pendidikan tinggi", "Tema ekonomi Islam": "Religiusitas, kepuasan finansial",
     "Temuan kunci": "Religiusitas, harapan, dan kepuasan finansial memengaruhi literasi mahasiswa",
     "Tantangan": "Generalisasi terbatas pada konteks kampus", "Relevansi RQ": "RQ1, RQ3"},
    {"Penulis dan tahun": "Albaity, M., & Rahman, M. (2019)", "Negara/konteks": "Uni Emirat Arab",
     "Tujuan": "Mengukur literasi keuangan syariah dan niat menggunakan bank syariah", "Metodologi": "Survei; pemodelan persamaan struktural (SEM)",
     "Populasi/sampel": "Responden dewasa (lihat artikel asli)", "Jenis literasi": "Literasi keuangan syariah",
     "Strategi/media pendidikan": "Tidak spesifik; fokus pada adopsi institusi", "Tema ekonomi Islam": "Perbankan syariah",
     "Temuan kunci": "Literasi keuangan syariah berkaitan dengan niat lebih besar mengadopsi perbankan syariah",
     "Tantangan": "Desain lintas-seksi membatasi inferensi kausal", "Relevansi RQ": "RQ2, RQ4"},
    {"Penulis dan tahun": "Ahmad, G. N., Widyastuti, U., Susanti, S., & Mukhibad, H. (2020)", "Negara/konteks": "Indonesia",
     "Tujuan": "Mengidentifikasi determinan literasi keuangan syariah", "Metodologi": "Survei kuantitatif",
     "Populasi/sampel": "Responden Muslim Indonesia (lihat artikel asli)", "Jenis literasi": "Literasi keuangan syariah",
     "Strategi/media pendidikan": "Faktor pendidikan dan sosiodemografis", "Tema ekonomi Islam": "Determinan literasi",
     "Temuan kunci": "Faktor sosiodemografis dan pendidikan memengaruhi tingkat literasi",
     "Tantangan": "Cakupan variabel determinan terbatas", "Relevansi RQ": "RQ1, RQ3"},
    {"Penulis dan tahun": "Solikhatun, I., & Sari, R. C. (2023)", "Negara/konteks": "Indonesia",
     "Tujuan": "Menguji edukasi keuangan syariah berbasis augmented reality (AR)", "Metodologi": "Kuasi-eksperimen",
     "Populasi/sampel": "Pelajar pada pembelajaran jarak jauh (lihat artikel asli)", "Jenis literasi": "Literasi keuangan syariah",
     "Strategi/media pendidikan": "Media augmented reality (AR); pembelajaran jarak jauh", "Tema ekonomi Islam": "Sosialisasi keuangan syariah",
     "Temuan kunci": "Media AR meningkatkan motivasi dan pengetahuan keuangan syariah",
     "Tantangan": "Skalabilitas dan ketersediaan infrastruktur", "Relevansi RQ": "RQ3"},
    {"Penulis dan tahun": "Kasri, R. A., & Yuniar, A. M. (2021)", "Negara/konteks": "Indonesia",
     "Tujuan": "Mengkaji determinan pembayaran zakat digital", "Metodologi": "Survei; kerangka UTAUT",
     "Populasi/sampel": "223 Muslim Indonesia pengguna platform daring", "Jenis literasi": "Literasi zakat (keuangan sosial Islam)",
     "Strategi/media pendidikan": "Platform pembayaran zakat digital", "Tema ekonomi Islam": "Zakat, keuangan sosial Islam",
     "Temuan kunci": "Literasi zakat meningkatkan niat membayar zakat melalui platform digital",
     "Tantangan": "Kepercayaan dan keamanan platform", "Relevansi RQ": "RQ3, RQ4"},
    {"Penulis dan tahun": "Anisa, S. N., & Fajri AF, M. S. (2025)", "Negara/konteks": "Indonesia",
     "Tujuan": "Menguji pengaruh literasi keuangan syariah terhadap keputusan investasi Gen Z", "Metodologi": "Survei kuantitatif (working paper)",
     "Populasi/sampel": "Generasi Z (lihat artikel asli)", "Jenis literasi": "Literasi keuangan syariah",
     "Strategi/media pendidikan": "Tidak spesifik; konteks pasar modal syariah", "Tema ekonomi Islam": "Investasi halal, pasar modal syariah",
     "Temuan kunci": "Literasi keuangan syariah berpengaruh positif dan signifikan terhadap keputusan investasi Gen Z",
     "Tantangan": "Status working paper; perlu telaah sejawat penuh", "Relevansi RQ": "RQ1, RQ4"},
]
fields = ["Penulis dan tahun", "Negara/konteks", "Tujuan", "Metodologi", "Populasi/sampel",
          "Jenis literasi", "Strategi/media pendidikan", "Tema ekonomi Islam", "Temuan kunci",
          "Tantangan", "Relevansi RQ"]
for idx, st in enumerate(studies, 1):
    h2("Studi %d. %s" % (idx, st["Penulis dan tahun"]))
    t = table(len(fields), 2)
    for i, f in enumerate(fields):
        cell(t.rows[i].cells[0], f, bold=True, size=9)
        cell(t.rows[i].cells[1], st[f], size=9)
    for r in t.rows:
        r.cells[0].width = Cm(4.5); r.cells[1].width = Cm(11.0)
    P(after=2)

# LAMPIRAN E
h1("Lampiran E. Diagram Alir PRISMA 2020")
if os.path.exists("prisma_flow_id.png"):
    p = P(align="center", after=2, before=4)
    p.add_run().add_picture("prisma_flow_id.png", width=Cm(14.0))
caption("Gambar E1. Diagram alir PRISMA 2020 untuk proses seleksi studi "
        "(jumlah bersifat ilustratif; difinalkan setelah ekstraksi lengkap).")

p = P(after=6, before=8)
run(p, "Catatan keterulangan: Seluruh sumber yang dikutip bersifat nyata dan dapat diverifikasi. Jumlah "
"rekaman pada Lampiran A dan diagram PRISMA bersifat ilustratif dan harus difinalkan dengan menjalankan "
"ulang string pencarian pada tanggal pengajuan. Daftar referensi lengkap juga disediakan dalam berkas RIS "
"terpisah untuk memudahkan verifikasi dan pengelolaan sitasi.", italic=True, size=9)

def assert_no_dash():
    bad = []
    for para in doc.paragraphs:
        if "\u2014" in para.text or "\u2013" in para.text:
            bad.append(para.text[:50])
    for tb in doc.tables:
        for row in tb.rows:
            for c in row.cells:
                if "\u2014" in c.text or "\u2013" in c.text:
                    bad.append(c.text[:50])
    return bad

bad = assert_no_dash()
if bad:
    raise SystemExit("DASH terlarang ditemukan: %r" % bad)

doc.save(OUT)
print("Saved:", OUT, "| tabel:", len(doc.tables), "| tanpa dash terlarang")
